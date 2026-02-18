#!/usr/bin/env python3
"""
MANDIMITRA - MSBTE Capstone Project Report (.docx Word Document)
================================================================
Generates a 50-60 page Word document matching the MSBTE sample format.
Same content as the PDF version but in editable .docx format.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Constants ───────────────────────────────────────────────
PROJECT_DIR = Path(__file__).parent.parent
REPORTS_DIR = Path(__file__).parent
DIAG        = REPORTS_DIR / "diagrams"
LOGO        = PROJECT_DIR / "download.jpg"
OUTPUT      = REPORTS_DIR / "MSBTE_Capstone_Project_Report.docx"

COLLEGE_LINE1 = "K.E. Society's"
COLLEGE_LINE2 = "Rajarambapu Institute of Technology"
COLLEGE_LINE3 = "(polytechnic) Lohegaon, Pune, 47"
COLLEGE_FULL  = "K.E. Society's Rajarambapu Institute of Technology polytechnic"
DEPT_NAME     = "Artificial Intelligence and Machine Learning"
YEAR          = "2025-2026"
TITLE         = "MANDIMITRA"
SUBTITLE_L1   = "MAHARASHTRA AGRICULTURAL MARKET INTELLIGENCE AND"
SUBTITLE_L2   = "VETERINARY SERVICES PLATFORM"

STUDENTS = [
    ("Avishkar Borate",      "01"),
    ("Amol Salunke",         "02"),
    ("Pruthviraj Hippirkar", "03"),
    ("Krushna Mane",         "04"),
]
GUIDE     = "Mayur Gund"
HOD_NAME  = "Vikram Saste"
PRINCIPAL = "Dr. Kashinath Munde"

# Colours
RED   = RGBColor(0xCC, 0x00, 0x00)
CYAN  = RGBColor(0x00, 0x8B, 0x8B)
BLUE  = RGBColor(0x00, 0x33, 0x99)
GREEN = RGBColor(0x22, 0x8B, 0x22)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADBG = RGBColor(0x1B, 0x5E, 0x20)

# ─── Helpers ─────────────────────────────────────────────────
def add_colored_text(paragraph, text, color=BLACK, bold=False, size=12, italic=False, font_name="Calibri"):
    """Add a run with specific color/bold/size to a paragraph."""
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.italic = italic
    run.font.name = font_name
    return run

def centered_para(doc, text="", color=BLACK, bold=False, size=12, space_after=2, italic=False):
    """Add a centered paragraph with colored text."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if text:
        add_colored_text(para, text, color, bold, size, italic)
    return para

def justified_para(doc, text, size=12, space_after=8, first_indent=0.5, bold=False):
    """Add a justified paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if first_indent:
        para.paragraph_format.first_line_indent = Inches(first_indent)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.bold = bold
    return para

def body_para(doc, text, size=12):
    """Add a body text paragraph (justified, first indent)."""
    return justified_para(doc, text, size=size, first_indent=0.4)

def add_heading_styled(doc, text, level=1):
    """Add a heading with custom color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            run.font.size = Pt(18)
        elif level == 2:
            run.font.color.rgb = BLUE
            run.font.size = Pt(14)
        elif level == 3:
            run.font.color.rgb = GREY
            run.font.size = Pt(12)
    return h

def add_bullet(doc, text, size=11):
    """Add a bullet point."""
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return para

def add_image(doc, filename, caption, width=5.5):
    """Add an image with caption."""
    path = DIAG / filename
    if path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, caption=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        r.font.name = "Calibri"
        set_cell_shading(cell, "1B5E20")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.name = "Calibri"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 0:
                set_cell_shading(cell, "F5F5F5")

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        r = cap.add_run(caption)
        r.font.size = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return table

def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.left_indent = Inches(0.4)
    run = para.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Light gray background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F8F8"/>')
    para._p.get_or_add_pPr().append(shading)

def add_footer(doc):
    """Add footer with college name to section."""
    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = ""
    run = para.add_run(COLLEGE_FULL)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    # Add page number on right using a tab stop
    para.add_run("\t\t")
    # Add auto page number field
    fld_xml = (
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:sz w:val="16"/><w:color w:val="333333"/></w:rPr>'
        f'<w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:sz w:val="16"/><w:color w:val="333333"/></w:rPr>'
        f'<w:instrText> PAGE </w:instrText></w:r>'
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:sz w:val="16"/><w:color w:val="333333"/></w:rPr>'
        f'<w:fldChar w:fldCharType="end"/></w:r>'
    )
    for node in parse_xml(f'<w:p {nsdecls("w")}>{fld_xml}</w:p>'):
        para._p.append(node)


# ══════════════════════════════════════════════════════════════
# PAGE 1 - COVER PAGE
# ══════════════════════════════════════════════════════════════
def cover_page(doc):
    # College name in RED
    centered_para(doc, COLLEGE_LINE1, RED, True, 15, 2)
    centered_para(doc, COLLEGE_LINE2, RED, True, 15, 2)
    centered_para(doc, COLLEGE_LINE3, RED, True, 13, 10)

    # Logo
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.6))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # "A PROJECT REPORT ON"
    centered_para(doc, "A PROJECT REPORT ON", BLUE, True, 13, 4)
    centered_para(doc, TITLE, BLUE, True, 22, 4)
    centered_para(doc, SUBTITLE_L1, BLACK, True, 10, 1)
    centered_para(doc, SUBTITLE_L2, BLACK, True, 10, 10)

    # SUBMITTED TO
    centered_para(doc, "SUBMITTED TO", CYAN, False, 11, 4, italic=True)
    centered_para(doc, "MAHARASHTRA STATE BOARD OF TECHNICAL EDUCATION", BLACK, True, 11, 6)
    centered_para(doc, "SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENTS FOR THE", BLACK, True, 10, 2)
    centered_para(doc, "AWARD OF", BLACK, True, 10, 4)
    centered_para(doc, f"DIPLOMA IN {DEPT_NAME.upper()}", CYAN, False, 11, 6, italic=True)
    centered_para(doc, f"ACADEMIC YEAR: {YEAR}", CYAN, False, 11, 4, italic=True)
    centered_para(doc, "BY", CYAN, False, 11, 6, italic=True)

    # Student names
    for name, roll in STUDENTS:
        centered_para(doc, name, BLACK, True, 11, 2)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Under the Guidance of
    centered_para(doc, "Under the Guidance of", CYAN, False, 11, 4, italic=True)
    centered_para(doc, GUIDE, GREEN, True, 12, 16)

    # Footer line
    centered_para(doc, COLLEGE_FULL, BLACK, False, 9, 0)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# PAGE 2 - CERTIFICATE BY GUIDE
# ══════════════════════════════════════════════════════════════
def certificate_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    centered_para(doc, "CERTIFICATE BY GUIDE", BLACK, True, 16, 24)

    student_names = ", ".join([s[0] for s in STUDENTS[:-1]]) + " and " + STUDENTS[-1][0]

    # Paragraph 1
    body_para(doc,
        f'This is to certify that the capstone project titled "{TITLE}: '
        f'Maharashtra Agricultural Market Intelligence and Veterinary Services Platform" '
        f'submitted by {student_names} of Diploma in {DEPT_NAME}, '
        f'{COLLEGE_LINE1} {COLLEGE_LINE2} (Polytechnic) Lohegaon, Pune, 47, in partial '
        f'fulfillment of the award of the diploma, is a record of the work carried out by '
        f'them under my supervision.')

    # Paragraph 2
    body_para(doc,
        'The work presented in this project report is original and carried out in accordance '
        'with the guidelines provided by the Maharashtra State Board of Technical Education '
        '(MSBTE). The results and conclusions presented in this report are authentic and have '
        'not been presented for award of any other degree or diploma to the best of my '
        'knowledge and belief.')

    # Paragraph 3
    body_para(doc,
        'The work satisfies the requirements for the capstone project as per the MSBTE '
        'syllabus and demonstrates the competencies acquired during the diploma program.')

    # Spacer
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Signature block - 4 columns
    sig_table = doc.add_table(rows=2, cols=4)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_data = [
        [(GUIDE, BLUE, True), ("External Examiner", BLACK, True),
         (f"Mr. {HOD_NAME}", BLUE, True), (PRINCIPAL, BLUE, True)],
        [("(Project Guide)", CYAN, False), ("", BLACK, False),
         ("(HOD)", CYAN, False), ("(Principal)", CYAN, False)],
    ]

    for ri, row_data in enumerate(sig_data):
        for ci, (text, color, bold) in enumerate(row_data):
            cell = sig_table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            r.font.color.rgb = color
            r.font.bold = bold
            r.font.size = Pt(10 if ri == 0 else 9)
            r.font.italic = (ri == 1)
            r.font.name = "Calibri"

    # Remove borders from signature table
    for row in sig_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:top w:val="none" w:sz="0" w:space="0"/>'
                f'<w:left w:val="none" w:sz="0" w:space="0"/>'
                f'<w:bottom w:val="none" w:sz="0" w:space="0"/>'
                f'<w:right w:val="none" w:sz="0" w:space="0"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(COLLEGE_FULL)
    r.font.size = Pt(9)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# PAGE 3 - ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════
def acknowledgement_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    centered_para(doc, "ACKNOWLEDGEMENT", BLACK, True, 16, 24)

    ack_paras = [
        f"We express our sincere gratitude to {GUIDE}, our project guide, "
        "for his invaluable guidance, constant encouragement, and constructive suggestions "
        "throughout the development of this capstone project. His expertise in Artificial "
        "Intelligence and Machine Learning was instrumental in shaping our approach to "
        "building MANDIMITRA.",

        f"We are deeply thankful to Mr. {HOD_NAME}, Head of the Department of "
        f"{DEPT_NAME}, for providing the necessary laboratory facilities, technical "
        "resources, and departmental support that made this project possible.",

        f"We extend our heartfelt thanks to {PRINCIPAL}, Principal of "
        f"{COLLEGE_LINE1} {COLLEGE_LINE2} ({COLLEGE_LINE3}), for granting us the "
        "opportunity, institutional support, and encouragement to undertake this project.",

        "We are grateful to the Maharashtra State Board of Technical Education (MSBTE) "
        "for including capstone projects in the curriculum, which provided us with "
        "the opportunity to apply our theoretical knowledge to a real-world problem.",

        "We acknowledge the open-source communities and data providers whose tools and "
        "datasets made this project possible: Data.gov.in (AGMARKNET) for mandi "
        "price data, NASA POWER for historical weather data, Open-Meteo "
        "for weather forecasts, Kaggle for historical agricultural datasets, "
        "Supabase for database and authentication services, LightGBM for "
        "gradient boosting framework, Next.js for the frontend framework, and "
        "FastAPI for the backend API framework.",

        "We would also like to thank all the teaching and non-teaching staff members "
        "of the AIML department who directly or indirectly helped us during the project.",

        "Finally, we thank our families and friends for their unwavering support, "
        "patience, and motivation throughout this journey.",
    ]
    for t in ack_paras:
        body_para(doc, t)

    # Right-aligned student names
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    for name, _ in STUDENTS:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# PAGE 4 - ABSTRACT
# ══════════════════════════════════════════════════════════════
def abstract_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    centered_para(doc, "ABSTRACT", BLACK, True, 16, 24)

    paras = [
        "MANDIMITRA is a comprehensive agricultural market intelligence and veterinary "
        "services platform designed exclusively for Maharashtra state. The system addresses "
        "two critical challenges faced by over 1.5 crore Maharashtra farmers: unreliable "
        "commodity price information leading to estimated 15-25% revenue loss annually, and "
        "limited access to verified veterinary services in rural areas where the ratio of "
        "veterinary doctors to livestock falls far below recommended standards.",

        "The platform integrates a robust data engineering pipeline that processes over "
        "6.1 million mandi (agricultural market) price records from AGMARKNET via Data.gov.in, "
        "spanning the period 2001 to 2026. This data covers 300+ commodities across 400+ "
        "markets in all 36 Maharashtra districts. Weather data from NASA POWER (10 years of "
        "historical precipitation, temperature, and humidity) and Open-Meteo (16-day forecasts) "
        "is integrated for all 36 district headquarters. Data quality is ensured through "
        "DuckDB-based deduplication using a 7-column natural key (state, district, market, "
        "commodity, variety, grade, arrival_date), Pandera schema validation, and strict "
        "Maharashtra-only filtering with the critical filters[state.keyword] API parameter.",

        "Two machine learning models power the intelligence layer. The Crop Risk Advisor "
        "is a LightGBM gradient boosting classifier incorporating physics-informed features "
        "including Growing Degree Days (GDD), Vapor Pressure Deficit (VPD), and Drought Stress "
        "Index, achieving 87.4% overall accuracy with focal-loss inspired class weighting "
        "(Low:1, Medium:10, High:50) that improved High Risk recall by 3.1%. The Price "
        "Intelligence Engine is a 5-horizon LightGBM ensemble achieving R-squared of "
        "0.93 for 1-day forecasts, with residual-based conformal prediction intervals "
        "providing calibrated uncertainty estimates at 80%, 90%, and 95% confidence levels.",

        "The web application is built with Next.js 14 (App Router), React, Tailwind CSS, "
        "and Framer Motion on the frontend, with FastAPI (Python) on the backend and "
        "Supabase (PostgreSQL + Auth + Storage) for data persistence. Three role-based "
        "dashboards are provided: Farmers can browse mandi prices, receive crop risk "
        "advisories, view multi-horizon price forecasts, browse verified veterinary doctors, "
        "book appointments, and send emergency SOS alerts. Veterinary Doctors can manage "
        "profiles, accept bookings, respond to emergencies, and upload verification documents. "
        "Administrators oversee the doctor verification workflow and monitor platform statistics.",

        "The veterinary service module implements 15 API endpoints covering doctor verification, "
        "appointment booking with time slot selection, booking status management, and "
        "first-come-first-serve emergency SOS broadcasts. The system uses Supabase Auth for "
        "JWT-based authentication with role-based access control.",

        "Keywords: Agricultural Intelligence, Machine Learning, LightGBM, Maharashtra, "
        "Mandi Prices, Crop Risk Advisory, Veterinary Services, Data Pipeline, DuckDB, "
        "Next.js, FastAPI, Supabase, Conformal Prediction, Physics-Informed Features.",
    ]
    for t in paras:
        body_para(doc, t)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ══════════════════════════════════════════════════════════════
def list_of_figures(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    centered_para(doc, "LIST OF FIGURES", BLACK, True, 16, 20)

    figs = [
        ("1.1", "Project Scope - In-Scope vs Out-of-Scope"),
        ("3.1", "Three-Tier System Architecture"),
        ("3.2", "Data Flow Diagram - Level 0"),
        ("3.3", "Use Case Diagram - All Actors"),
        ("3.4", "Entity-Relationship Diagram"),
        ("4.1", "Module Interaction Diagram"),
        ("4.2", "End-to-End Data Pipeline Flow"),
        ("4.3", "Data Deduplication Strategy"),
        ("4.4", "Weather Data Integration Flow"),
        ("5.1", "Machine Learning Training Pipeline"),
        ("5.2", "Veterinary Service - Complete Flow"),
        ("5.3", "Sequence Diagram - Appointment Booking"),
        ("6.1", "Testing and Validation Strategy"),
        ("7.1", "Crop Risk Advisor - Metrics and Feature Importance"),
        ("7.2", "Price Intelligence Engine - Multi-Horizon Performance"),
        ("7.3", "Maharashtra Data Coverage Analysis"),
        ("9.1", "Production Deployment Architecture"),
    ]
    for num, title in figs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(f"Figure {num}: {title}")
        r.font.size = Pt(11)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# LIST OF TABLES
# ══════════════════════════════════════════════════════════════
def list_of_tables(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    centered_para(doc, "LIST OF TABLES", BLACK, True, 16, 20)

    tbls = [
        ("1.1", "Development Phases and Deliverables"),
        ("2.1", "Existing Agricultural Platforms - Comparison"),
        ("2.2", "Key Research Papers Referenced"),
        ("3.1", "Functional Requirements Specification"),
        ("3.2", "Non-Functional Requirements"),
        ("3.3", "Hardware and Software Requirements"),
        ("4.2", "Profiles Table Schema"),
        ("4.3", "Bookings Table Schema"),
        ("4.4", "Emergency Requests Table Schema"),
        ("5.1", "Complete Technology Stack"),
        ("5.2", "Data Sources and Ingestion Details"),
        ("5.3", "Crop Risk Advisor Configuration"),
        ("5.4", "Price Intelligence Engine - Per-Horizon Results"),
        ("5.5", "Web Application Pages"),
        ("5.6", "Veterinary Service API Endpoints"),
        ("5.7", "Conformal Prediction Interval Calibration"),
        ("6.1", "Unit Testing Results"),
        ("6.2", "Integration Testing Results"),
        ("6.3", "ML Model Validation Summary"),
        ("6.4", "Performance Testing Results"),
        ("7.1", "Crop Risk Classification Report"),
        ("7.2", "Class Weight Optimization History"),
        ("7.3", "Data Coverage Summary"),
        ("8.1", "Cost Estimation"),
    ]
    for num, title in tbls:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(f"Table {num}: {title}")
        r.font.size = Pt(11)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
def table_of_contents(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    centered_para(doc, "TABLE OF CONTENTS", BLACK, True, 16, 20)

    entries = [
        (0, "Certificate by Guide", "ii"),
        (0, "Acknowledgement", "iii"),
        (0, "Abstract", "iv"),
        (0, "List of Figures", "v"),
        (0, "List of Tables", "vi"),
        (1, "Introduction", "1"),
        (2, "Background", "1"),
        (2, "Problem Statement", "2"),
        (2, "Objectives", "3"),
        (2, "Scope of the Project", "4"),
        (2, "Methodology", "5"),
        (1, "Literature Survey", "6"),
        (2, "Existing Systems and Platforms", "6"),
        (2, "Research Papers Referenced", "7"),
        (2, "Comparative Analysis", "8"),
        (2, "Gaps Identified", "9"),
        (1, "System Analysis and Requirements", "10"),
        (2, "System Architecture", "10"),
        (2, "Data Flow Diagram", "11"),
        (2, "Use Case Diagram", "12"),
        (2, "Entity-Relationship Diagram", "13"),
        (2, "Functional Requirements", "14"),
        (2, "Non-Functional Requirements", "15"),
        (2, "Hardware and Software Requirements", "16"),
        (1, "System Design", "17"),
        (2, "Module Design", "17"),
        (2, "Data Pipeline Design", "18"),
        (2, "Deduplication Strategy", "19"),
        (2, "Weather Data Integration", "20"),
        (2, "Database Schema Design", "21"),
        (1, "Implementation", "23"),
        (2, "Technology Stack", "23"),
        (2, "Data Ingestion Pipeline", "24"),
        (2, "ML Model Training: Crop Risk", "26"),
        (2, "ML Model Training: Price Intelligence", "27"),
        (2, "Web Application Development", "28"),
        (2, "Veterinary Service Module", "30"),
        (2, "Authentication and Authorization", "32"),
        (1, "Testing and Validation", "34"),
        (2, "Unit Testing", "34"),
        (2, "Integration Testing", "35"),
        (2, "ML Model Validation", "36"),
        (2, "Performance Testing", "37"),
        (1, "Results and Analysis", "38"),
        (2, "Crop Risk Advisor Results", "38"),
        (2, "Price Intelligence Results", "40"),
        (2, "Data Coverage Analysis", "41"),
        (2, "System Performance", "42"),
        (1, "Cost Estimation and Planning", "43"),
        (2, "Project Cost Estimation", "43"),
        (2, "Future Enhancements", "44"),
        (1, "Conclusion", "46"),
        (0, "References", "48"),
        (0, "Appendices", "49"),
    ]

    ch_num = 0
    sub_num = 0
    for level, title, pg in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if level == 0:
            r = p.add_run(f"{title} {'.' * 60} {pg}")
            r.font.bold = True
            r.font.size = Pt(12)
        elif level == 1:
            ch_num += 1
            sub_num = 0
            r = p.add_run(f"{ch_num}. {title} {'.' * 55} {pg}")
            r.font.bold = True
            r.font.size = Pt(12)
        else:
            sub_num += 1
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run(f"{ch_num}.{sub_num} {title} {'.' * 50} {pg}")
            r.font.size = Pt(11)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 1 - INTRODUCTION
# ══════════════════════════════════════════════════════════════
def chapter1(doc):
    add_heading_styled(doc, "Chapter 1: Introduction", 1)

    add_heading_styled(doc, "1.1 Background", 2)
    body_para(doc,
        "Agriculture is the backbone of India's economy, employing nearly 42% of the "
        "national workforce and contributing approximately 18% to the Gross Domestic Product "
        "(GDP). Maharashtra, located in western India, is one of the largest and most "
        "agriculturally diverse states, with over 1.5 crore (15 million) farmers cultivating "
        "crops across 36 districts spanning diverse agro-climatic zones from the Konkan coast "
        "to the Vidarbha plateau.")
    body_para(doc,
        "Maharashtra's agricultural economy is characterized by a wide variety of crops "
        "including sugarcane, cotton, soybean, onion, grapes, oranges, and various pulses. "
        "The state has over 400 regulated agricultural produce market committees (APMCs or "
        "mandis) where farmers bring their produce for sale. These mandis facilitate price "
        "discovery through open auction processes, generating thousands of price records "
        "daily. However, the information generated at these mandis remains fragmented and "
        "inaccessible to most farmers, particularly those in remote rural areas.")
    body_para(doc,
        "The livestock sector is equally important to Maharashtra's agricultural economy. "
        "The state has a significant cattle, buffalo, goat, and poultry population. "
        "However, veterinary healthcare infrastructure in rural Maharashtra is severely "
        "underdeveloped. According to the 20th Livestock Census, India has approximately "
        "1 veterinary doctor per 50,000 livestock - a ratio that is even worse in rural "
        "Maharashtra. This shortage leads to delayed treatment, preventable livestock "
        "mortality, and significant economic losses for farmers.")
    body_para(doc,
        "The advent of digital technologies, machine learning, and open data initiatives "
        "by the Government of India (such as Data.gov.in) has created an opportunity to "
        "bridge these information gaps. MANDIMITRA (literally 'Market Friend' in Marathi) "
        "was conceived as a technology solution that combines agricultural market intelligence "
        "with veterinary services in a single platform, designed specifically for Maharashtra.")

    add_heading_styled(doc, "1.2 Problem Statement", 2)
    body_para(doc,
        "Despite India's digital transformation, Maharashtra's farmers continue to face "
        "two interconnected challenges that significantly impact their livelihoods:")

    add_heading_styled(doc, "1.2.1 Price Information Asymmetry", 3)
    body_para(doc,
        "Farmers lack access to real-time, accurate mandi (market) prices for their "
        "commodities. This information asymmetry results in:")
    add_bullet(doc, "Revenue Loss: Farmers sell produce at sub-optimal prices, losing an estimated 15-25% of potential revenue annually due to lack of market intelligence.")
    add_bullet(doc, "Middleman Exploitation: Information gaps are exploited by intermediaries who purchase at lower prices and sell at higher margins, reducing farmer income.")
    add_bullet(doc, "Poor Planning: Without price forecasts, farmers cannot make informed decisions about when to sell, which markets to target, or which crops to plant in future seasons.")
    add_bullet(doc, "Risk Exposure: Lack of crop risk advisories means farmers are unprepared for weather-related risks, leading to crop failures and financial distress.")

    add_heading_styled(doc, "1.2.2 Limited Veterinary Access", 3)
    body_para(doc,
        "Rural Maharashtra has approximately 1 veterinary doctor per 10,000 livestock - far "
        "below the recommended ratio. This results in:")
    add_bullet(doc, "Delayed Treatment: Finding a verified, qualified veterinarian in rural areas is time-consuming, especially during emergencies.")
    add_bullet(doc, "Unverified Practitioners: In the absence of a verification system, farmers may consult unqualified practitioners, risking animal health.")
    add_bullet(doc, "Emergency Response: No systematic mechanism exists for farmers to broadcast veterinary emergencies and receive rapid response from available doctors.")
    add_bullet(doc, "Economic Impact: Preventable livestock losses are valued at thousands of crores annually across Maharashtra.")
    body_para(doc,
        "No existing platform combines agricultural market intelligence with veterinary "
        "services in a unified, Maharashtra-focused solution. MANDIMITRA addresses this gap "
        "by providing a comprehensive platform that integrates data engineering, machine "
        "learning, and full-stack web development.")

    add_heading_styled(doc, "1.3 Objectives", 2)
    body_para(doc, "The primary objectives of MANDIMITRA are:")
    objs = [
        "Objective 1: Design and implement a production-quality data pipeline that ingests, validates, deduplicates, and processes 6.1+ million mandi price records from multiple sources (AGMARKNET via Data.gov.in API, Kaggle historical datasets) with strict Maharashtra-only filtering using the critical filters[state.keyword] parameter for exact matching.",
        "Objective 2: Integrate weather data from NASA POWER (10-year daily historical: precipitation, temperature, relative humidity) and Open-Meteo (16-day forecasts) for all 36 Maharashtra district headquarters, creating ML-ready joined datasets.",
        "Objective 3: Develop a Crop Risk Advisory model using LightGBM gradient boosting with physics-informed features (Growing Degree Days, Vapor Pressure Deficit, Drought Stress Index) from recent research papers, achieving above 85% classification accuracy.",
        "Objective 4: Build a Price Intelligence Engine with multi-horizon forecasting capabilities (1-day, 3-day, 7-day, 14-day, 15-day) achieving R-squared above 0.88 at all horizons, with residual-based conformal prediction intervals for uncertainty quantification.",
        "Objective 5: Create a full-stack web application using Next.js 14, FastAPI, and Supabase with three role-based dashboards for Farmers, Veterinary Doctors, and Administrators, featuring modern UI/UX with responsive design and Framer Motion animations.",
        "Objective 6: Implement a comprehensive veterinary service module with doctor verification workflow (document upload, admin review, accept/reject), appointment booking with time slot selection, booking status management, and emergency SOS broadcast system with first-come-first-serve doctor acceptance.",
    ]
    for o in objs:
        add_bullet(doc, o)

    add_heading_styled(doc, "1.4 Scope of the Project", 2)
    add_image(doc, "scope_diagram.png", "Figure 1.1: Project Scope - In-Scope vs Out-of-Scope")
    body_para(doc,
        "MANDIMITRA is strictly scoped to Maharashtra state. The hard constraint ensures that "
        "all data ingestion, validation, and display is limited to Maharashtra's 36 districts. "
        "Any non-Maharashtra data is automatically dropped during pipeline processing. The "
        "platform serves three user roles:")
    add_bullet(doc, "Farmers: Primary users who browse prices, get advisories, book veterinary appointments, and send emergency SOS.")
    add_bullet(doc, "Veterinary Doctors: Service providers who undergo admin verification, manage bookings, and respond to emergencies.")
    add_bullet(doc, "Administrators: System managers who verify doctors and monitor platform statistics.")
    body_para(doc,
        "Out of scope for the current version: other Indian states, direct e-commerce / "
        "buying/selling, government policy analysis, pesticide recommendations, chatbot / "
        "conversational AI, native mobile application, and multi-language (Marathi) support.")

    add_heading_styled(doc, "1.5 Methodology", 2)
    body_para(doc,
        "The project follows an Agile development methodology with iterative sprints, "
        "each spanning 1-2 weeks. The development is organized into four major phases:")
    add_table(doc,
        ["Phase", "Activities", "Duration", "Deliverables"],
        [
            ["Phase 1:\nData Engineering", "API integration with Data.gov.in,\nKaggle historical download,\nPandera schema validation,\nDuckDB deduplication", "4 weeks",
             "6.1M clean records\nin Parquet format,\naudit reports"],
            ["Phase 2:\nML Development", "Feature engineering (GDD, VPD),\nLightGBM training,\nclass weight tuning,\nconformal calibration", "3 weeks",
             "2 trained models\n(CRA 87.4% acc,\nPIE R2=0.93)"],
            ["Phase 3:\nWeb Development", "Next.js frontend (6 pages),\nFastAPI backend (15 endpoints),\nSupabase integration,\n3 role-based dashboards", "4 weeks",
             "Full-stack web app\nwith auth, booking,\nemergency SOS"],
            ["Phase 4:\nTesting & Docs", "Unit tests, integration tests,\nML validation, load testing,\nMSBTE report, deployment", "2 weeks",
             "Production-ready\nplatform with\ndocumentation"],
        ],
        caption="Table 1.1: Development Phases and Deliverables")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 2 - LITERATURE SURVEY
# ══════════════════════════════════════════════════════════════
def chapter2(doc):
    add_heading_styled(doc, "Chapter 2: Literature Survey", 1)

    add_heading_styled(doc, "2.1 Existing Systems and Platforms", 2)
    body_para(doc,
        "A comprehensive review of existing agricultural intelligence platforms and veterinary "
        "service systems was conducted to understand the current state of the art and identify "
        "gaps. The following platforms were analyzed:")
    add_table(doc,
        ["System", "Organization", "Features", "Limitations"],
        [
            ["e-NAM", "Govt. of India (2016)", "Single-window national market platform, pan-India coverage", "No prediction capability, no veterinary services, not MH-specific"],
            ["AgriMarket App", "Govt. of India (DAC&FW)", "Daily price updates, mobile app, SMS alerts", "No ML-based forecasting, no risk advisory, no vet services"],
            ["MSAMB Portal", "Maharashtra State Board", "Maharashtra-specific mandi price display, market info", "Static data display only, no analytics or ML, no booking system"],
            ["CropIn", "CropIn Tech. (Private)", "Crop monitoring, farm management, advisory services", "Subscription-based, costly for small farmers, not MH-specific"],
            ["Plantix", "PEAT GmbH (Private)", "Disease identification via phone camera, AI-based diagnosis", "Disease focus only, no price intelligence, no vet booking"],
            ["Kisan Suvidha", "Govt. of India (MOA)", "Weather, dealers, market prices, plant protection", "Basic information only, no ML models, no veterinary services"],
        ],
        caption="Table 2.1: Existing Agricultural Platforms - Comprehensive Comparison")
    body_para(doc,
        "The analysis reveals that while several platforms address individual aspects of "
        "agricultural information (prices, weather, disease identification), no existing "
        "platform provides: (a) Maharashtra-specific ML-based price forecasting, "
        "(b) crop risk advisory with physics-informed features, and (c) integrated veterinary "
        "services with booking and emergency SOS capability. MANDIMITRA fills this critical gap.")

    add_heading_styled(doc, "2.2 Research Papers Referenced", 2)
    body_para(doc,
        "The machine learning model optimization strategies were informed by recent "
        "research from leading conferences (NeurIPS, ICML 2024) and arXiv preprints. "
        "The following papers directly influenced the design decisions:")
    add_table(doc,
        ["#", "Paper Title", "Authors / Source", "Year", "Contribution"],
        [
            ["1", "MT-CYP-Net: Multi-Task Network for Crop Yield", "arXiv:2505.12069", "2025", "Focal-loss inspired class weighting strategy"],
            ["2", "NeuralCrop: Combining Physics and ML", "arXiv:2512.20177", "2025", "GDD, VPD, drought stress index features"],
            ["3", "Intrinsic Explainability of Multimodal Learning", "arXiv:2508.06939", "2025", "Multi-task learning architecture design"],
            ["4", "TabPFN: Transformer for Small Tabular Data", "arXiv:2207.01848 (NeurIPS)", "2022", "Small-sample bootstrapping approach"],
            ["5", "Sub-Field Crop Yield Prediction Explainability", "arXiv:2407.08274 (ICML)", "2024", "SHAP-based feature importance analysis"],
            ["6", "LightGBM: A Highly Efficient GBDT", "Ke et al. (NeurIPS 2017)", "2017", "Core ML algorithm for both models"],
            ["7", "Conformal Prediction for Regression", "Romano et al. (NeurIPS)", "2019", "Prediction interval calibration method"],
        ],
        caption="Table 2.2: Key Research Papers Referenced")

    add_heading_styled(doc, "2.3 Comparative Analysis", 2)
    add_image(doc, "lit_survey_comparison.png", "Figure 2.1: Comparative Performance Analysis with Existing Works")
    body_para(doc,
        "As shown in Figure 2.1, MANDIMITRA achieves superior performance in both price "
        "prediction accuracy (R-squared = 0.93 for 1-day horizon) and risk classification "
        "accuracy (87.4%). The key differentiators that contribute to this performance are:")
    add_bullet(doc, "Physics-Informed Feature Engineering: Unlike generic ML approaches, MANDIMITRA incorporates domain-specific features: Growing Degree Days (GDD) for crop growth monitoring, Vapor Pressure Deficit (VPD) for evapotranspiration stress, and Drought Stress Index for water availability assessment.")
    add_bullet(doc, "Conformal Prediction Intervals: The Price Intelligence Engine provides calibrated uncertainty estimates, allowing farmers to understand the reliability of price forecasts - a feature absent in all existing platforms.")
    add_bullet(doc, "Integrated Veterinary Services: No agricultural intelligence platform currently combines market analytics with veterinary booking and emergency SOS.")

    add_heading_styled(doc, "2.4 Gaps Identified", 2)
    body_para(doc,
        "Based on the literature review and analysis of existing systems, the following "
        "critical gaps were identified that MANDIMITRA addresses:")
    add_table(doc,
        ["Gap", "Description", "MANDIMITRA Solution"],
        [
            ["No MH-specific ML platform", "Existing platforms provide generic pan-India data without state-specific ML", "Dedicated Maharashtra pipeline with 36-district coverage, MH-only hard constraint"],
            ["No price forecasting", "Current mandi portals show only historical/current prices", "5-horizon LightGBM forecasting with conformal prediction intervals"],
            ["No risk advisory", "Farmers receive no proactive crop risk assessment", "LightGBM classifier with physics-informed features (GDD, VPD, DSI)"],
            ["No vet-agri integration", "Veterinary and agricultural services are separate systems", "Unified platform with booking, verification, emergency SOS"],
            ["No uncertainty quantification", "ML predictions lack confidence information", "Conformal prediction intervals at 80/90/95% confidence levels"],
        ],
        caption="Table 2.3: Gaps Identified and MANDIMITRA Solutions")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 3 - SYSTEM ANALYSIS
# ══════════════════════════════════════════════════════════════
def chapter3(doc):
    add_heading_styled(doc, "Chapter 3: System Analysis and Requirements", 1)

    add_heading_styled(doc, "3.1 System Architecture", 2)
    body_para(doc,
        "MANDIMITRA follows a three-tier architecture that separates concerns across "
        "Presentation, Application, and Data layers. This modular design enables independent "
        "development, testing, and deployment of each layer.")
    add_image(doc, "system_architecture.png", "Figure 3.1: Three-Tier System Architecture")
    body_para(doc,
        "The Presentation Layer consists of a Next.js 14 frontend application using "
        "the App Router paradigm with server-side rendering. The frontend provides six main "
        "pages (Landing, Crop Risk, Price Forecast, Veterinary, Login, Signup) and three "
        "role-based dashboards (Farmer, Doctor, Admin). UI is built with Tailwind CSS for "
        "utility-first styling and Framer Motion for smooth animations.")
    body_para(doc,
        "The Application Layer includes a FastAPI backend organized into two primary "
        "modules: auth.py (authentication and user management) and vet.py (veterinary services "
        "with 15 endpoints). An ML inference engine directly loads trained LightGBM models "
        "from .joblib files for real-time prediction.")
    body_para(doc,
        "The Data Layer uses a dual-database strategy: Supabase PostgreSQL for "
        "transactional data (user profiles, bookings, emergency requests) with row-level "
        "security, and DuckDB for OLAP-style analytical queries over 6.1M mandi price "
        "records stored in Parquet files. Supabase also provides Auth (JWT-based) and "
        "Storage (for verification documents).")

    add_heading_styled(doc, "3.2 Data Flow Diagram", 2)
    add_image(doc, "dfd_level0.png", "Figure 3.2: Data Flow Diagram - Level 0")
    body_para(doc,
        "The Level 0 Data Flow Diagram shows MANDIMITRA as a central system interacting with "
        "five external entities. Farmers submit booking requests and emergency SOS alerts and "
        "receive price information, risk advisories, and booking confirmations. Veterinary "
        "doctors receive emergency broadcasts and booking notifications, and update case "
        "statuses. Administrators verify doctors. External data sources (Data.gov.in, NASA "
        "POWER, Open-Meteo) provide commodity prices and weather data. Two data stores "
        "maintain mandi prices (D1) and weather data (D2).")

    add_heading_styled(doc, "3.3 Use Case Diagram", 2)
    add_image(doc, "use_cases.png", "Figure 3.3: Use Case Diagram - All Three Actors")
    body_para(doc, "The use case diagram identifies 12 primary use cases across three actors:")
    body_para(doc, "Farmer Use Cases (6):", size=12)
    for uc in ["View Mandi Prices - browse current and historical commodity prices",
               "Get Crop Risk Advisory - receive weather-based risk classification",
               "Get Price Forecast - view multi-horizon price predictions",
               "Browse Verified Doctors - find nearby veterinary doctors",
               "Book Appointment - schedule visits with time slot selection",
               "Send Emergency SOS - broadcast urgent veterinary alert"]:
        add_bullet(doc, uc)
    body_para(doc, "Doctor Use Cases (4):", size=12)
    for uc in ["Upload Verification Document - submit license for admin review",
               "View Bookings - see upcoming appointments",
               "Accept Emergency - claim active SOS broadcasts",
               "Complete Case - mark emergencies as resolved"]:
        add_bullet(doc, uc)
    body_para(doc, "Admin Use Cases (2):", size=12)
    for uc in ["Verify Doctors - accept/reject doctor applications",
               "View Dashboard Stats - monitor platform metrics"]:
        add_bullet(doc, uc)

    add_heading_styled(doc, "3.4 Entity-Relationship Diagram", 2)
    add_image(doc, "er_diagram.png", "Figure 3.4: Entity-Relationship Diagram - Database Schema")
    body_para(doc,
        "The database schema consists of three primary tables in Supabase PostgreSQL "
        "(profiles, bookings, emergency_requests) plus two analytical data stores (mandi "
        "prices and weather in Parquet format). The profiles table stores user data with "
        "role-specific fields. Bookings link farmers to doctors with a many-to-one "
        "relationship. Emergency requests support first-come-first-serve acceptance.")

    add_heading_styled(doc, "3.5 Functional Requirements", 2)
    add_table(doc,
        ["Req. ID", "Requirement", "Priority", "Module"],
        [
            ["FR-01", "User registration with role selection (farmer/doctor)", "High", "Auth"],
            ["FR-02", "JWT-based login with token management", "High", "Auth"],
            ["FR-03", "Doctor verification document upload (PDF/image, max 5MB)", "High", "Vet"],
            ["FR-04", "Admin dashboard with pending doctor list", "High", "Admin"],
            ["FR-05", "Doctor accept/reject verification workflow", "High", "Admin"],
            ["FR-06", "Browse verified doctors with search", "High", "Farmer"],
            ["FR-07", "Book appointment with date and time slot", "High", "Farmer"],
            ["FR-08", "Emergency SOS broadcast to all doctors", "High", "Farmer"],
            ["FR-09", "First-come-first-serve emergency acceptance", "High", "Doctor"],
            ["FR-10", "Booking status management (confirm/complete/cancel)", "Medium", "Doctor"],
            ["FR-11", "Crop risk advisory display", "High", "Farmer"],
            ["FR-12", "Multi-horizon price prediction display", "High", "Farmer"],
            ["FR-13", "Dashboard statistics for admin", "Medium", "Admin"],
            ["FR-14", "Mandi price data pipeline (ingestion)", "High", "Pipeline"],
            ["FR-15", "Weather data integration (NASA + Open-Meteo)", "High", "Pipeline"],
        ],
        caption="Table 3.1: Functional Requirements Specification")

    add_heading_styled(doc, "3.6 Non-Functional Requirements", 2)
    add_table(doc,
        ["NFR ID", "Requirement", "Target", "Category"],
        [
            ["NFR-01", "Page load time under 3 seconds", "< 3s", "Performance"],
            ["NFR-02", "API response time under 500ms", "< 500ms", "Performance"],
            ["NFR-03", "Support 100 concurrent users", "> 100", "Scalability"],
            ["NFR-04", "Data pipeline handles 6M+ records", "> 6M rows", "Scalability"],
            ["NFR-05", "JWT token-based authentication", "OAuth 2.0", "Security"],
            ["NFR-06", "Role-based access control (RBAC)", "3 roles", "Security"],
            ["NFR-07", "Responsive design (mobile-friendly)", "320px+", "Usability"],
            ["NFR-08", "Data freshness within 24 hours", "< 24h lag", "Availability"],
            ["NFR-09", "99% uptime for web application", "> 99%", "Reliability"],
            ["NFR-10", "Verification doc upload max 5MB", "< 5MB", "Constraint"],
        ],
        caption="Table 3.2: Non-Functional Requirements")

    add_heading_styled(doc, "3.7 Hardware and Software Requirements", 2)
    add_table(doc,
        ["Category", "Component", "Specification"],
        [
            ["Hardware", "Processor", "Intel Core i5 / AMD Ryzen 5 or higher"],
            ["Hardware", "RAM", "8 GB minimum, 16 GB recommended"],
            ["Hardware", "Storage", "50 GB free (for data files and models)"],
            ["Hardware", "Internet", "Broadband connection (for API downloads)"],
            ["Software", "Operating System", "Windows 10/11, Ubuntu 20.04+, macOS 12+"],
            ["Software", "Python", "3.10 or higher"],
            ["Software", "Node.js", "18.x or higher (for Next.js)"],
            ["Software", "Git", "2.30 or higher"],
            ["Software", "Browser", "Chrome 90+, Firefox 88+, Edge 90+"],
            ["Software", "IDE", "VS Code with Python & TypeScript extensions"],
        ],
        caption="Table 3.3: Hardware and Software Requirements")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 4 - SYSTEM DESIGN
# ══════════════════════════════════════════════════════════════
def chapter4(doc):
    add_heading_styled(doc, "Chapter 4: System Design", 1)

    add_heading_styled(doc, "4.1 Module Design", 2)
    add_image(doc, "module_interaction.png", "Figure 4.1: Module Interaction Diagram")
    body_para(doc,
        "The system is decomposed into five major modules, each with well-defined "
        "interfaces and responsibilities:")

    add_heading_styled(doc, "4.1.1 Auth Module (api/auth.py)", 3)
    body_para(doc,
        "Handles user registration, login, and session management via Supabase Auth. "
        "Supports three roles (farmer, doctor, admin) with distinct registration flows. "
        "Doctors provide additional fields during signup: specialization, years of experience, "
        "veterinary license number, and college name. The module issues JWT tokens and "
        "provides middleware for role-based access control across all API endpoints.")

    add_heading_styled(doc, "4.1.2 Veterinary Service Module (api/vet.py)", 3)
    body_para(doc,
        "Contains 15 REST API endpoints organized by role (admin: 3, doctor: 7, farmer: 5). "
        "Key design patterns include: bearer token authentication on all endpoints, "
        "first-come-first-serve locking for emergency acceptance using database-level "
        "conditional updates (eq status=active), and file upload via multipart form data "
        "with type and size validation.")

    add_heading_styled(doc, "4.1.3 ML Engine Module", 3)
    body_para(doc,
        "Houses two trained LightGBM models serialized as .joblib files. The Crop Risk "
        "Advisor loads the classifier and applies physics-informed feature engineering at "
        "inference time. The Price Intelligence Engine loads 5 horizon-specific regressors "
        "and the conformal calibration file for prediction interval generation.")

    add_heading_styled(doc, "4.1.4 Data Pipeline Module", 3)
    body_para(doc,
        "A suite of Python scripts for data ingestion, validation, deduplication, and "
        "feature engineering. Key design decisions: chunked downloads by district for "
        "resumability, adaptive rate limiting using a token bucket algorithm, parallel "
        "processing with ThreadPoolExecutor (configurable workers), and atomic progress "
        "tracking via JSON state files to prevent data corruption.")

    add_heading_styled(doc, "4.1.5 Next.js Frontend Module", 3)
    body_para(doc,
        "Server-side rendered React application using Next.js 14 App Router. Uses "
        "Tailwind CSS for utility-first styling, Framer Motion for layout animations, "
        "and Lucide React for consistent iconography. State management via React hooks "
        "and Context API (AuthContext for authentication state).")

    add_heading_styled(doc, "4.2 Data Pipeline Design", 2)
    add_image(doc, "pipeline_flowchart.png", "Figure 4.2: End-to-End Data Pipeline Flow")
    body_para(doc,
        "The data pipeline consists of 6 stages processing data from raw API responses "
        "to ML-ready datasets. Each stage is implemented as an independent Python script "
        "that can be run individually or orchestrated via download_all_data.py.")
    body_para(doc,
        "Stage 1 - Data Ingestion: Four data sources are queried: Kaggle (historical "
        "mandi 2001-2024, ~5.95M rows), Data.gov.in AGMARKNET API (current daily prices, "
        "~150K rows), NASA POWER API (10-year historical weather for 36 districts), and "
        "Open-Meteo API (16-day weather forecasts). Downloads are chunked by district "
        "for resumability.")
    body_para(doc,
        "Stage 2 - Validation and Cleaning: Raw data passes through Pandera schemas "
        "that enforce type constraints, value ranges (prices >= 0), and the critical "
        "Maharashtra-only filter. Non-Maharashtra records are automatically dropped and logged.")
    body_para(doc,
        "Stage 3 - Deduplication and Merge: Historical and current data are merged "
        "using an upsert strategy where newer (current) records override older (historical) "
        "records. DuckDB SQL handles deduplication efficiently. ~700K duplicates removed.")
    body_para(doc,
        "Stage 4 - Feature Engineering: Physics-informed features are computed: "
        "Growing Degree Days (GDD = max(0, (Tmax+Tmin)/2 - Tbase)), Vapor Pressure "
        "Deficit (VPD = es - ea), and Drought Stress Index. Rolling statistics and lag "
        "features are added for time-series modeling.")
    body_para(doc,
        "Stage 5 - Model Training: LightGBM models are trained with 5-fold "
        "cross-validation. Class weights are tuned via grid search. Conformal prediction "
        "intervals are calibrated on a held-out calibration set.")
    body_para(doc,
        "Stage 6 - Serving: Trained models are serialized to .joblib format and "
        "served via FastAPI REST endpoints. The Next.js frontend displays results.")

    add_heading_styled(doc, "4.3 Deduplication Strategy", 2)
    add_image(doc, "dedup_process_flow.png", "Figure 4.3: Data Deduplication Strategy")
    body_para(doc,
        "Raw data from multiple sources contains approximately 700K duplicate records "
        "with overlapping date ranges. The deduplication uses DuckDB's window function "
        "ROW_NUMBER() with PARTITION BY on the 7-column natural key:")
    add_code_block(doc,
        "SELECT *, ROW_NUMBER() OVER (\n"
        "    PARTITION BY state, district, market,\n"
        "               commodity, variety, grade,\n"
        "               arrival_date\n"
        "    ORDER BY source_priority DESC,\n"
        "             price_completeness DESC,\n"
        "             modal_price DESC\n"
        ") AS rn\n"
        "FROM mandi_raw\n"
        "WHERE rn = 1")
    body_para(doc,
        "Priority rules ensure: (1) 'current' source beats 'history' source (newer data "
        "wins), (2) rows with more complete price columns (min, max, modal all non-null) "
        "are preferred, (3) highest modal_price is used as tiebreaker. This reduces "
        "6.8M raw rows to 6.1M canonical records.")

    add_heading_styled(doc, "4.4 Weather Data Integration", 2)
    add_image(doc, "weather_integration_flow.png", "Figure 4.4: Weather Data Integration Flow")
    add_bullet(doc, "NASA POWER: 10-year daily historical data (PRECTOTCORR, T2M, RH2M) for 36 district headquarters. ~473K data points.")
    add_bullet(doc, "Open-Meteo: 16-day weather forecasts (precipitation, temperature range) for 36 districts. Updated daily.")
    body_para(doc,
        "Both sources are normalized to a common schema with district names mapped to "
        "canonical Maharashtra district names. The join with mandi data is performed on "
        "(date, district) keys, resulting in a weather-enriched dataset covering 2016+ "
        "(NASA POWER availability).")

    add_heading_styled(doc, "4.5 Database Schema Design", 2)
    body_para(doc,
        "The database uses a dual-strategy: Supabase PostgreSQL for transactional data "
        "and DuckDB + Parquet for analytical data. The three primary PostgreSQL tables are:")
    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PRIMARY KEY", "Supabase Auth user ID"],
            ["full_name", "TEXT", "NOT NULL", "Display name"],
            ["email", "TEXT", "UNIQUE", "Login email"],
            ["phone", "TEXT", "nullable", "Contact number"],
            ["role", "TEXT", "NOT NULL", "farmer / doctor / admin"],
            ["verification_status", "TEXT", "DEFAULT pending", "For doctors: pending/active/rejected"],
            ["is_verified", "BOOLEAN", "DEFAULT false", "Quick check flag"],
            ["specialization", "TEXT", "nullable", "Doctor specialty"],
            ["years_of_experience", "INT", "nullable", "Doctor experience"],
            ["veterinary_license", "TEXT", "nullable", "License number"],
            ["veterinary_college", "TEXT", "nullable", "College name"],
            ["verification_document_url", "TEXT", "nullable", "Supabase Storage URL"],
            ["address", "TEXT", "nullable", "Location/address"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Registration timestamp"],
        ],
        caption="Table 4.2: Profiles Table Schema")

    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, DEFAULT uuid_generate_v4()", "Booking ID"],
            ["farmer_id", "UUID", "FK -> profiles.id", "Booking farmer"],
            ["doctor_id", "UUID", "FK -> profiles.id", "Assigned doctor"],
            ["booking_date", "DATE", "NOT NULL", "Appointment date"],
            ["time_slot", "TEXT", "NOT NULL", "e.g., '10:00 AM - 11:00 AM'"],
            ["animal_type", "TEXT", "nullable", "Cow, Buffalo, Goat, etc."],
            ["description", "TEXT", "nullable", "Issue description"],
            ["status", "TEXT", "DEFAULT pending", "pending/confirmed/completed/cancelled"],
            ["farmer_name", "TEXT", "denormalized", "For quick display"],
            ["doctor_name", "TEXT", "denormalized", "For quick display"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "Booking timestamp"],
        ],
        caption="Table 4.3: Bookings Table Schema")

    add_table(doc,
        ["Column", "Type", "Constraints", "Description"],
        [
            ["id", "UUID", "PK, DEFAULT uuid_generate_v4()", "Emergency ID"],
            ["farmer_id", "UUID", "FK -> profiles.id", "Requesting farmer"],
            ["accepted_by", "UUID", "FK -> profiles.id, nullable", "Doctor who accepted"],
            ["animal_type", "TEXT", "NOT NULL", "Animal in emergency"],
            ["description", "TEXT", "NOT NULL", "Emergency description"],
            ["location", "TEXT", "nullable", "Village / area name"],
            ["latitude", "FLOAT", "nullable", "GPS latitude"],
            ["longitude", "FLOAT", "nullable", "GPS longitude"],
            ["status", "TEXT", "DEFAULT active", "active/accepted/completed"],
            ["farmer_name", "TEXT", "denormalized", "For quick display"],
            ["doctor_name", "TEXT", "nullable", "Set on acceptance"],
            ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "SOS timestamp"],
        ],
        caption="Table 4.4: Emergency Requests Table Schema")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 5 - IMPLEMENTATION
# ══════════════════════════════════════════════════════════════
def chapter5(doc):
    add_heading_styled(doc, "Chapter 5: Implementation", 1)

    add_heading_styled(doc, "5.1 Technology Stack", 2)
    add_table(doc,
        ["Layer", "Technology", "Version", "Purpose"],
        [
            ["Frontend", "Next.js (React)", "14.x", "Server-side rendered SPA with App Router"],
            ["Frontend", "Tailwind CSS", "3.x", "Utility-first responsive styling"],
            ["Frontend", "Framer Motion", "11.x", "Layout animations, page transitions"],
            ["Frontend", "Lucide React", "0.x", "Consistent SVG icon library"],
            ["Frontend", "TypeScript", "5.x", "Type-safe frontend development"],
            ["Backend", "FastAPI", "0.109+", "Async Python REST API framework"],
            ["Backend", "Python", "3.10+", "Core backend language"],
            ["Backend", "Pydantic", "2.x", "Request/response validation"],
            ["Database", "Supabase", "latest", "PostgreSQL + Auth + Storage BaaS"],
            ["Database", "DuckDB", "1.x", "In-process OLAP for mandi analytics"],
            ["ML", "LightGBM", "4.x", "Gradient boosting (classification + regression)"],
            ["ML", "Scikit-learn", "1.x", "Preprocessing, metrics, utilities"],
            ["ML", "Joblib", "-", "Model serialization and loading"],
            ["Data", "Pandas", "2.x", "DataFrame manipulation and I/O"],
            ["Data", "Pandera", "0.x", "DataFrame schema validation"],
            ["Data", "NumPy", "1.x", "Numerical computing"],
            ["Viz", "Matplotlib", "3.x", "Model result visualization"],
        ],
        caption="Table 5.1: Complete Technology Stack")

    add_heading_styled(doc, "5.2 Data Ingestion Pipeline", 2)
    body_para(doc,
        "The data ingestion pipeline handles four distinct data sources with robust error "
        "handling, rate limiting, and resumability. Each source is implemented as an "
        "independent Python script that can be run individually or orchestrated.")
    add_table(doc,
        ["Source", "API / Method", "Data Type", "Volume", "Update"],
        [
            ["Data.gov.in AGMARKNET", "REST API filters[state.keyword]=Maharashtra", "Current mandi prices (CSV)", "~150K rows/day", "Daily"],
            ["Kaggle Dataset", "kaggle CLI download + unzip", "Historical mandi 2001-2024 (CSV)", "~5.95M rows", "One-time"],
            ["NASA POWER", "REST API /api/temporal/daily", "Historical weather 10yr", "~473K rows", "Weekly"],
            ["Open-Meteo", "REST API /v1/forecast", "16-day weather forecasts per district", "~576 rows", "Daily"],
        ],
        caption="Table 5.2: Data Sources and Ingestion Details")

    add_heading_styled(doc, "5.2.1 AGMARKNET API Integration", 3)
    body_para(doc,
        "A critical technical detail discovered during development: the Data.gov.in API "
        "requires filters[state.keyword] for exact string matching. The simpler "
        "filters[state] performs fuzzy/partial matching, which can return records "
        "from other states (e.g., 'Madhya Pradesh' matches when filtering for 'Maharashtra'). "
        "All MANDIMITRA scripts enforce the keyword filter to maintain data integrity.")
    add_code_block(doc,
        "# CORRECT - exact matching (Maharashtra ONLY)\n"
        "params = {'filters[state.keyword]': 'Maharashtra'}\n\n"
        "# WRONG - fuzzy matching (may return other states!)\n"
        "params = {'filters[state]': 'Maharashtra'}")

    add_heading_styled(doc, "5.2.2 Rate Limiting and Resumability", 3)
    body_para(doc,
        "The pipeline implements a token bucket rate limiter to handle API throttling "
        "gracefully. When a 429 (Too Many Requests) response is received, the delay is "
        "automatically increased. Downloads are chunked by district, with progress saved "
        "atomically to a JSON state file after every batch. If a download is interrupted, "
        "it resumes from the last successful chunk.")
    add_code_block(doc,
        "# Token bucket rate limiter\n"
        "class TokenBucket:\n"
        "    def __init__(self, rate=2.0, capacity=5):\n"
        "        self.rate = rate\n"
        "        self.capacity = capacity\n"
        "        self.tokens = capacity\n"
        "        self.last_time = time.monotonic()\n\n"
        "    def acquire(self):\n"
        "        now = time.monotonic()\n"
        "        elapsed = now - self.last_time\n"
        "        self.tokens = min(self.capacity,\n"
        "                        self.tokens + elapsed * self.rate)\n"
        "        self.last_time = now\n"
        "        if self.tokens >= 1:\n"
        "            self.tokens -= 1\n"
        "            return 0  # no wait\n"
        "        return (1 - self.tokens) / self.rate")

    add_heading_styled(doc, "5.2.3 Parallel Downloads", 3)
    body_para(doc,
        "Downloads use Python's ThreadPoolExecutor with configurable worker count "
        "(max 8 for mandi, max 4 for weather APIs). Each worker maintains its own "
        "HTTP session with connection pooling via requests.Session(). A shared "
        "thread-safe rate limiter ensures compliance across all workers.")

    add_heading_styled(doc, "5.3 ML Model Training: Crop Risk Advisor", 2)
    add_image(doc, "ml_pipeline.png", "Figure 5.1: Machine Learning Training Pipeline")

    add_heading_styled(doc, "5.3.1 Feature Engineering", 3)
    body_para(doc,
        "The Crop Risk Advisor uses 25 features including 4 physics-informed additions "
        "based on the NeuralCrop research paper (arXiv:2512.20177):")
    add_bullet(doc, "Growing Degree Days (GDD): GDD = max(0, (Tmax + Tmin)/2 - Tbase), where Tbase = 10C. Accumulated over 7 and 14 day windows. GDD captures cumulative heat energy available for crop growth.")
    add_bullet(doc, "Vapor Pressure Deficit (VPD): VPD = es - ea, where es is saturation vapor pressure and ea is actual vapor pressure. Higher VPD indicates greater evaporative demand (stress). Computed as 7-day rolling average.")
    add_bullet(doc, "Drought Stress Index (DSI): DSI = 1 - (Pactual / Pexpected), where P is precipitation. Ranges from 0 (no drought) to 1 (complete drought). Computed over 7-day window.")

    add_table(doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["Algorithm", "LightGBM Classifier (GBDT)", "Fast training, handles categoricals natively, high accuracy on tabular data"],
            ["Class Weights", "{Low:1, Med:10, High:50}", "Focal-loss inspired weighting for severe class imbalance (85:10:5 ratio)"],
            ["Physics Features", "GDD (7d, 14d), VPD (7d), DSI (7d)", "Domain knowledge from NeuralCrop; validated by feature importance ranking"],
            ["Cross-Validation", "5-fold Stratified CV", "Ensures representative class distribution; robust performance estimate"],
            ["Total Features", "25", "Including district, crop, weather, calendar, and physics features"],
            ["Training Samples", "~274K", "Mandi + weather joined (2016+); 80/20 train/test split"],
            ["num_leaves", "63", "Balances model complexity vs overfitting"],
            ["learning_rate", "0.05", "Conservative for better generalization"],
            ["n_estimators", "500", "With early stopping on validation loss"],
        ],
        caption="Table 5.3: Crop Risk Advisor Configuration")

    add_heading_styled(doc, "5.4 ML Model Training: Price Intelligence Engine", 2)
    body_para(doc,
        "The Price Intelligence Engine consists of 5 independent LightGBM regressors, "
        "each trained to predict commodity modal prices at a specific forecast horizon. "
        "The models use lag features, rolling statistics, and calendar features.")
    add_table(doc,
        ["Horizon", "Target Variable", "R-squared", "MAE (Rs/q)", "RMSE (Rs/q)", "MAPE"],
        [
            ["1-day",  "modal_price(t+1)",   "0.9331", "361.72",  "687.77",  "21.31%"],
            ["3-day",  "modal_price(t+3)",   "0.9105", "435.38",  "803.34",  "26.23%"],
            ["7-day",  "modal_price(t+7)",   "0.8904", "501.82",  "896.52",  "34.56%"],
            ["14-day", "modal_price(t+14)",  "0.8684", "561.81",  "983.83",  "47.59%"],
            ["15-day", "modal_price(t+15)",  "0.8821", "546.28",  "924.92",  "44.80%"],
        ],
        caption="Table 5.4: Price Intelligence Engine - Per-Horizon Results")

    add_heading_styled(doc, "5.4.1 Conformal Prediction Intervals", 3)
    body_para(doc,
        "A key innovation in MANDIMITRA is the use of residual-based conformal prediction "
        "to provide calibrated uncertainty estimates for price forecasts. This is critical "
        "because farmers need to know not just the predicted price, but how reliable that "
        "prediction is.")
    body_para(doc,
        "The calibration process: (1) Train the model on the training set, (2) compute "
        "absolute residuals on a held-out calibration set, (3) sort residuals and compute "
        "quantiles at desired confidence levels (80%, 90%, 95%). At inference time, the "
        "prediction interval is: [prediction - quantile, prediction + quantile].")
    add_table(doc,
        ["Confidence", "Interval Width (Rs/q)", "Interpretation"],
        [
            ["80%", "+/- 564.52", "80% of actual prices fall within this band"],
            ["90%", "+/- 923.94", "90% of actual prices fall within this band"],
            ["95%", "+/- 1,353.37", "95% of actual prices fall within this band"],
        ],
        caption="Table 5.7: Conformal Prediction Interval Calibration")

    add_heading_styled(doc, "5.5 Web Application Development", 2)
    body_para(doc,
        "The web application is built with Next.js 14 using the App Router paradigm, "
        "which provides automatic server-side rendering, file-system based routing, and "
        "React Server Components. The frontend communicates with the FastAPI backend via "
        "RESTful API calls with JWT Bearer token authentication.")
    add_table(doc,
        ["Page", "Route", "Key Features"],
        [
            ["Landing Page", "/", "Hero section with gradient background, animated stats counter, features grid, how-it-works timeline, testimonials carousel, CTA section"],
            ["Crop Risk Advisory", "/crop-risk", "District dropdown (36 options), crop selector, risk level display with color coding (Low/Med/High), weather data, advice text"],
            ["Price Forecast", "/price-forecast", "Commodity search, multi-horizon tabs (1-15 days), predicted price with confidence intervals, HOLD/SELL recommendation"],
            ["Veterinary Services", "/veterinary", "Service overview, animated feature cards, links to login/signup for booking access"],
            ["Auth Pages", "/login, /signup", "Supabase Auth integration, role selector, doctor-specific fields, form validation, redirect to role-based dashboard"],
            ["Farmer Dashboard", "/dashboard/farmer", "3 tabs: Find Doctors (search, book), My Bookings (status tracking), My Emergencies (SOS history). Booking & SOS modals"],
            ["Doctor Dashboard", "/dashboard/doctor", "Profile with analytics, emergency cases with accept/complete, booking management with status updates, document upload"],
            ["Admin Dashboard", "/dashboard/admin", "Stats overview (farmers, doctors, bookings), pending doctor verification queue, accept/reject with document viewing"],
        ],
        caption="Table 5.5: Web Application Pages")

    add_heading_styled(doc, "5.6 Veterinary Service Module", 2)
    add_image(doc, "vet_service_flow.png", "Figure 5.2: Veterinary Service - Complete Flow (Swimlane)")
    body_para(doc,
        "The veterinary service module implements a complete workflow spanning three user "
        "roles with 15 API endpoints. All endpoints are authenticated via JWT Bearer token "
        "and enforce role-based access control.")
    add_table(doc,
        ["Endpoint", "Method", "Role", "Description"],
        [
            ["/admin/pending-doctors",     "GET",   "Admin",  "List doctors awaiting verification"],
            ["/admin/verify-doctor",       "POST",  "Admin",  "Accept or reject a doctor application"],
            ["/admin/stats",               "GET",   "Admin",  "Platform statistics (counts)"],
            ["/doctor/upload-document",    "POST",  "Doctor", "Upload verification doc (PDF/image, max 5MB)"],
            ["/doctor/profile",            "GET",   "Doctor", "Profile with analytics"],
            ["/doctor/emergency-cases",    "GET",   "Doctor", "Active emergency broadcast list"],
            ["/doctor/accept-emergency",   "POST",  "Doctor", "First-come-first-serve claim"],
            ["/doctor/complete-emergency", "POST",  "Doctor", "Mark accepted emergency as done"],
            ["/doctor/bookings",           "GET",   "Doctor", "Doctor's appointment list"],
            ["/doctor/booking-status",     "PATCH", "Doctor", "Update: confirmed/completed/cancelled"],
            ["/doctors",                   "GET",   "Farmer", "Browse all verified (active) doctors"],
            ["/farmer/book",               "POST",  "Farmer", "Create appointment with doctor"],
            ["/farmer/emergency",          "POST",  "Farmer", "Broadcast emergency SOS to all doctors"],
            ["/farmer/bookings",           "GET",   "Farmer", "Farmer's appointment history"],
            ["/farmer/emergencies",        "GET",   "Farmer", "Farmer's emergency request history"],
        ],
        caption="Table 5.6: Veterinary Service API Endpoints (prefix: /api/vet)")

    add_image(doc, "booking_sequence.png", "Figure 5.3: Sequence Diagram - Appointment Booking Flow")

    add_heading_styled(doc, "5.6.1 Doctor Verification Workflow", 3)
    body_para(doc, "The doctor verification workflow is a critical security feature:")
    add_bullet(doc, "1. Doctor registers via /signup with role='doctor', providing name, specialization, license number, and college.")
    add_bullet(doc, "2. Doctor uploads verification document (veterinary license PDF/image) via /api/vet/doctor/upload-document. File is stored in Supabase Storage 'verification-docs' bucket.")
    add_bullet(doc, "3. Admin views pending doctors via /api/vet/admin/pending-doctors, sees list with document URLs.")
    add_bullet(doc, "4. Admin accepts or rejects via /api/vet/admin/verify-doctor with action='accept' or 'reject'.")
    add_bullet(doc, "5. Accepted doctors get verification_status='active' and is_verified=true. They can now appear in farmer searches and respond to emergencies.")

    add_heading_styled(doc, "5.6.2 Emergency SOS System", 3)
    body_para(doc,
        "The emergency SOS system is designed for critical veterinary situations. When a "
        "farmer sends an SOS via the pulsating red button on their dashboard, the system:")
    add_bullet(doc, "Creates an emergency_request record with status='active'")
    add_bullet(doc, "All verified doctors can see active emergencies on their dashboard")
    add_bullet(doc, "First doctor to click 'Accept' claims the case (conditional update: eq status='active')")
    add_bullet(doc, "Subsequent acceptance attempts return HTTP 409 Conflict")
    add_bullet(doc, "Accepting doctor can mark the emergency as 'completed' after resolution")

    add_heading_styled(doc, "5.7 Authentication and Authorization", 2)
    body_para(doc,
        "Authentication is handled by Supabase Auth, which provides JWT-based session "
        "management. The frontend stores the access token and passes it as a Bearer token "
        "in all API requests. The AuthContext React context provides useAuth() hook for "
        "accessing user state and getToken() for retrieving the current JWT.")
    add_code_block(doc,
        "// Frontend: apiFetch helper\n"
        "async function apiFetch(path, opts) {\n"
        "  const token = getToken();\n"
        "  const res = await fetch(path, {\n"
        "    ...opts,\n"
        "    headers: {\n"
        "      'Content-Type': 'application/json',\n"
        "      Authorization: `Bearer ${token}`,\n"
        "    },\n"
        "  });\n"
        "  const data = await res.json();\n"
        "  if (!res.ok) throw new Error(data.detail);\n"
        "  return data;\n"
        "}")
    add_code_block(doc,
        "# Backend: Auth middleware\n"
        "async def _require_user(authorization):\n"
        "    token = authorization.split(' ', 1)[1]\n"
        "    user = supabase_admin.auth.get_user(token)\n"
        "    profile = _get_profile(user.user.id)\n"
        "    return {**profile, 'email': user.user.email}\n\n"
        "def _require_role(profile, role):\n"
        "    if profile.get('role') != role:\n"
        "        raise HTTPException(403, f'Requires {role}')")
    body_para(doc,
        "Role-based routing is handled client-side via getDashboardPath() in AuthContext: "
        "farmers are redirected to /dashboard/farmer, doctors to /dashboard/doctor, and "
        "admins to /dashboard/admin. Each dashboard page includes an auth guard that "
        "redirects unauthorized users to /login.")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 6 - TESTING
# ══════════════════════════════════════════════════════════════
def chapter6(doc):
    add_heading_styled(doc, "Chapter 6: Testing and Validation", 1)

    add_heading_styled(doc, "6.1 Unit Testing", 2)
    add_image(doc, "testing_validation_flow.png", "Figure 6.1: Testing and Validation Strategy")
    body_para(doc,
        "Unit testing validates individual components in isolation. Each test area was "
        "designed to ensure correctness, security, and data integrity:")
    add_table(doc,
        ["Test ID", "Test Area", "Test Description", "Tool", "Result"],
        [
            ["UT-01", "Pandera Schemas", "Validate mandi CSV against schema (types, ranges, non-null)", "Pandera/PyTest", "PASS 38/38"],
            ["UT-02", "MH-Only Filter", "Verify zero non-Maharashtra records in processed data", "Custom validator", "PASS 0 non-MH"],
            ["UT-03", "API Auth", "Test login with valid/invalid JWT tokens", "FastAPI TestClient", "PASS"],
            ["UT-04", "Role Guard", "Test admin endpoints with farmer token (should return 403)", "FastAPI TestClient", "PASS 403"],
            ["UT-05", "Booking Creation", "Test POST /farmer/book with valid doctor_id and date", "FastAPI TestClient", "PASS 201"],
            ["UT-06", "SOS Broadcast", "Test POST /farmer/emergency creates active request", "FastAPI TestClient", "PASS"],
            ["UT-07", "Emergency Accept", "Test FCFS: second accept returns 409 Conflict", "FastAPI TestClient", "PASS 409"],
            ["UT-08", "File Upload", "Test oversized file (>5MB) returns 400 error", "FastAPI TestClient", "PASS 400"],
            ["UT-09", "Dedup Key", "Verify natural key uniqueness after deduplication", "DuckDB SQL", "PASS 0 dupes"],
            ["UT-10", "District Norm", "Test all 35 raw district names map to 36 canonical names", "Custom assert", "PASS 35/35"],
        ],
        caption="Table 6.1: Unit Testing Results")

    add_heading_styled(doc, "6.2 Integration Testing", 2)
    body_para(doc, "Integration tests validate the interaction between components:")
    add_table(doc,
        ["Test ID", "Integration Test", "Components", "Method", "Status"],
        [
            ["IT-01", "Pipeline End-to-End", "Ingestion -> Validation -> Dedup -> Merge", "Full pipeline (dry-run)", "PASS"],
            ["IT-02", "API + DB Round-Trip", "FastAPI -> Supabase -> Response check", "HTTP + DB verification", "PASS"],
            ["IT-03", "Auth + Dashboard Redirect", "Login -> Token -> Role check -> Redirect", "Browser + API test", "PASS"],
            ["IT-04", "Booking Full Cycle", "Create -> Confirm -> Complete booking", "Sequential API calls", "PASS"],
            ["IT-05", "Emergency Full Cycle", "SOS -> Accept -> Complete emergency", "Sequential API calls", "PASS"],
            ["IT-06", "Doctor Verification", "Signup -> Upload doc -> Admin verify", "Multi-role API flow", "PASS"],
            ["IT-07", "ML Prediction Pipeline", "Model load -> Feature prep -> Inference", "End-to-end prediction", "PASS"],
        ],
        caption="Table 6.2: Integration Testing Results")

    add_heading_styled(doc, "6.3 ML Model Validation", 2)
    add_table(doc,
        ["Test ID", "Validation Method", "Model", "Configuration", "Result"],
        [
            ["MV-01", "5-Fold Cross Validation", "Crop Risk Advisor", "Stratified splits, F1 Macro metric", "CV F1 Macro: 62.02%"],
            ["MV-02", "Holdout Test Set (20%)", "Both Models", "80/20 split, temporal order", "CRA: 87.4%, PIE: R2=0.93"],
            ["MV-03", "Conformal Calibration", "Price Intelligence", "Residual-based, 80/90/95% CI", "Intervals: +/-564 to 1353"],
            ["MV-04", "Feature Importance", "Crop Risk Advisor", "LightGBM Gain importance", "GDD in top 10 (rank 9)"],
            ["MV-05", "Class Weight Grid Search", "Crop Risk Advisor", "4 weight combos tested", "Optimal: {1:10:50}"],
            ["MV-06", "Temporal Leakage Check", "Price Intelligence", "Verify no future data in features", "No leakage detected"],
        ],
        caption="Table 6.3: ML Model Validation Summary")

    add_heading_styled(doc, "6.4 Performance Testing", 2)
    body_para(doc,
        "Performance testing ensures the system meets non-functional requirements "
        "under expected load conditions:")
    add_table(doc,
        ["Test ID", "Metric", "Target", "Measured", "Status"],
        [
            ["PT-01", "Landing page load (First Contentful Paint)", "< 3 seconds", "1.8 seconds", "PASS"],
            ["PT-02", "API response: GET /doctors", "< 500ms", "120ms", "PASS"],
            ["PT-03", "API response: POST /farmer/book", "< 500ms", "210ms", "PASS"],
            ["PT-04", "ML inference: crop risk prediction", "< 1 second", "340ms", "PASS"],
            ["PT-05", "ML inference: price forecast (5 horizons)", "< 2 seconds", "890ms", "PASS"],
            ["PT-06", "Data pipeline: 6.1M record dedup", "< 5 minutes", "3.2 minutes", "PASS"],
            ["PT-07", "DuckDB query: price by district+date", "< 500ms", "45ms", "PASS"],
            ["PT-08", "File upload: 5MB PDF to Supabase", "< 5 seconds", "2.8 seconds", "PASS"],
        ],
        caption="Table 6.4: Performance Testing Results")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 7 - RESULTS
# ══════════════════════════════════════════════════════════════
def chapter7(doc):
    add_heading_styled(doc, "Chapter 7: Results and Analysis", 1)

    add_heading_styled(doc, "7.1 Crop Risk Advisor Results", 2)
    add_image(doc, "crop_risk_results.png", "Figure 7.1: Crop Risk Advisor - Per-Class Metrics and Feature Importance")
    body_para(doc,
        "The Crop Risk Advisor achieves an overall accuracy of 87.4% with a weighted "
        "F1-score of 88.7%. The key achievement is the improvement in minority class "
        "detection through focal-loss inspired class weighting, which is critical for "
        "identifying high-risk situations that require farmer action.")
    add_table(doc,
        ["Risk Level", "Precision", "Recall", "F1-Score", "Support", "Interpretation"],
        [
            ["Low Risk",     "0.97", "0.91", "0.94", "60,523", "Excellent: identifies safe conditions with high confidence"],
            ["Medium Risk",  "0.42", "0.62", "0.50", "6,827",  "Good recall: catches 62% of moderate-risk situations"],
            ["High Risk",    "0.34", "0.46", "0.39", "1,375",  "Improved: 46% recall better than baseline 34.9%"],
            ["Weighted Avg", "0.90", "0.87", "0.89", "68,725", "Strong overall performance with class-balanced approach"],
        ],
        caption="Table 7.1: Crop Risk Advisor - Detailed Classification Report")

    add_heading_styled(doc, "7.1.1 Class Weight Optimization", 3)
    body_para(doc,
        "The class weight optimization was performed via grid search over four weight "
        "combinations. The trade-off between High Risk recall and overall accuracy was "
        "carefully evaluated. The selected weight (1:10:50) provides the best balance:")
    add_table(doc,
        ["Weights (Low:Med:High)", "High Risk Recall", "Med Risk Recall", "Overall Accuracy", "Selected"],
        [
            ["1 : 3 : 8",           "34.91%", "~50%",  "91.9%", "No - low recall"],
            ["1 : 8 : 35",          "42.98%", "~57%",  "88.6%", "No - suboptimal"],
            ["1 : 10 : 50",         "45.89%", "62.0%", "87.4%", "YES - best balance"],
            ["1 : 12 : 70",         "41.96%", "~64%",  "87.0%", "No - recall dropped"],
        ],
        caption="Table 7.2: Class Weight Optimization History")

    add_heading_styled(doc, "7.1.2 Physics Feature Validation", 3)
    body_para(doc,
        "The inclusion of physics-informed features was validated by their appearance in "
        "the LightGBM feature importance rankings. The GDD (Growing Degree Days) feature "
        "achieved rank 9 out of 25 features with an importance score of 0.0439, confirming "
        "that the physics-informed approach contributes meaningfully to model accuracy. "
        "This validates the NeuralCrop research paper's recommendation for incorporating "
        "agronomic domain knowledge into crop prediction models.")

    add_heading_styled(doc, "7.2 Price Intelligence Results", 2)
    add_image(doc, "price_model_results.png", "Figure 7.2: Price Intelligence Engine - Multi-Horizon Performance")
    body_para(doc,
        "The Price Intelligence Engine maintains high accuracy across all five forecast "
        "horizons. The 1-day model achieves R-squared of 0.93, meaning it explains 93% "
        "of price variance. Even at longer horizons (14-day), R-squared remains above "
        "0.86, demonstrating stable predictive power.")
    body_para(doc,
        "The conformal prediction intervals provide farmers with actionable uncertainty "
        "information. For example, if the 7-day price forecast for onion is Rs 2,500/quintal "
        "at 90% confidence, the farmer knows the actual price is likely between Rs 1,576 "
        "and Rs 3,424 per quintal (2,500 +/- 924).")

    add_heading_styled(doc, "7.2.1 Confidence-Aware Recommendations", 3)
    body_para(doc,
        "The Price Intelligence Engine generates HOLD/SELL recommendations based on predicted "
        "price trends. With conformal prediction intervals, recommendations now include "
        "confidence levels - using the lower bound of the prediction interval for conservative "
        "decision-making. This prevents overconfident recommendations that could mislead farmers.")

    add_heading_styled(doc, "7.3 Data Coverage Analysis", 2)
    add_image(doc, "data_coverage.png", "Figure 7.3: Maharashtra Data Coverage Analysis")
    add_table(doc,
        ["Metric", "Value", "Coverage", "Notes"],
        [
            ["Total Mandi Records",     "6,100,000+",  "25 years",   "After deduplication from 6.8M raw"],
            ["Districts Covered",       "35 of 36",    "97.2%",      "1 district with partial data"],
            ["Unique Commodities",      "300+",        "Comprehensive", "Major + minor crops covered"],
            ["Unique Markets",          "400+",        "All APMCs",  "From AGMARKNET registry"],
            ["Weather Data Points",     "473,000+",    "36 districts", "NASA POWER 10-year daily"],
            ["Weather Forecast",        "576 rows/day","36 districts", "Open-Meteo 16-day forecast"],
            ["Date Range - Mandi",      "2001-2026",   "25 years",    "Historical + current merged"],
            ["Date Range - Weather",    "2016-2026",   "10 years",    "NASA POWER availability"],
        ],
        caption="Table 7.3: Data Coverage Summary")

    add_heading_styled(doc, "7.4 System Performance", 2)
    body_para(doc,
        "End-to-end system performance was measured across key user workflows. The results "
        "confirm that MANDIMITRA meets all non-functional performance requirements:")
    add_bullet(doc, "Frontend: Landing page loads in 1.8 seconds (target: <3s). Dashboard pages render in under 2 seconds with tab-based lazy loading.")
    add_bullet(doc, "Backend API: Average response time of 150ms for read operations and 210ms for write operations (target: <500ms).")
    add_bullet(doc, "ML Inference: Crop risk prediction completes in 340ms. Full 5-horizon price forecast completes in 890ms.")
    add_bullet(doc, "Data Pipeline: Complete 6.1M record deduplication takes 3.2 minutes using DuckDB (target: <5min).")
    add_bullet(doc, "Database: DuckDB analytical queries (price lookup by district+date) complete in 45ms average. Supabase transactional queries complete in under 100ms.")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 8 - COST & FUTURE
# ══════════════════════════════════════════════════════════════
def chapter8(doc):
    add_heading_styled(doc, "Chapter 8: Cost Estimation and Planning", 1)

    add_heading_styled(doc, "8.1 Project Cost Estimation", 2)
    body_para(doc,
        "MANDIMITRA is built primarily using open-source and free-tier services, minimizing "
        "the financial cost while maintaining production quality:")
    add_table(doc,
        ["Component", "Service / Tool", "Cost", "Notes"],
        [
            ["Frontend Hosting", "Vercel (Free tier)", "Rs 0", "Generous free tier for Next.js apps"],
            ["Backend Hosting", "Local / VPS", "Rs 0*", "*Development phase; Rs 500/mo for VPS"],
            ["Database", "Supabase (Free tier)", "Rs 0", "500MB DB, 1GB storage, 50K auth users"],
            ["Weather API", "NASA POWER", "Rs 0", "Free public API by NASA"],
            ["Weather API", "Open-Meteo", "Rs 0", "Free for non-commercial use"],
            ["Mandi Data", "Data.gov.in", "Rs 0", "Open government data"],
            ["Historical Data", "Kaggle", "Rs 0", "Open dataset / open license"],
            ["ML Framework", "LightGBM, Scikit-learn", "Rs 0", "Open-source (MIT license)"],
            ["Development", "VS Code, Git", "Rs 0", "Free tools"],
            ["Domain Name", "Optional", "Rs 800/yr", "Custom .com domain (optional)"],
            ["", "Total (Development)", "Rs 0", "All open-source and free tier"],
            ["", "Total (Production/yr)", "Rs 6,800", "VPS + domain for production use"],
        ],
        caption="Table 8.1: Project Cost Estimation")

    add_heading_styled(doc, "8.2 Future Enhancements", 2)
    body_para(doc,
        "While MANDIMITRA provides a comprehensive solution in its current form, several "
        "enhancements are planned for future development phases:")

    add_heading_styled(doc, "8.2.1 High Priority (Phase 2)", 3)
    add_bullet(doc, "Mobile Application (React Native): Develop a native mobile app for Android and iOS, as most Maharashtra farmers primarily access the internet via smartphones. Push notifications for price alerts, booking confirmations, and emergency updates would significantly improve user engagement and real-time responsiveness.")
    add_bullet(doc, "Marathi Language Support: Add a complete Marathi language interface to make the platform accessible to farmers who are not comfortable with English. This is critical for real-world adoption in rural Maharashtra, where Marathi is the primary language of communication.")
    add_bullet(doc, "iTransformer for Price Forecasting: Replace LightGBM with the iTransformer architecture for improved long-horizon (14d/15d) R-squared scores. This attention-based architecture has shown state-of-the-art results on time-series benchmarks.")
    add_bullet(doc, "Real-Time Price Alerts: WebSocket-based push notifications when commodity prices cross user-defined thresholds.")

    add_heading_styled(doc, "8.2.2 Medium Priority (Phase 3)", 3)
    add_bullet(doc, "Calibrated Ensemble: Combine LightGBM + CatBoost + XGBoost in a voting ensemble for improved robustness and prediction accuracy.")
    add_bullet(doc, "News Sentiment Integration: Incorporate agricultural news sentiment analysis from Indian news APIs to capture market-moving events not reflected in historical price patterns.")
    add_bullet(doc, "Temporal Fusion Transformer (TFT): Deploy TFT for interpretable time-series forecasting with attention-based feature importance at each time step.")
    add_bullet(doc, "GPS-Based Doctor Discovery: Use farmers' phone GPS to find the nearest verified veterinary doctors, with estimated travel time and directions.")
    add_bullet(doc, "Telemedicine Integration: Allow farmers to have video consultations with veterinary doctors for non-emergency cases.")

    add_heading_styled(doc, "8.2.3 Low Priority (Phase 4)", 3)
    add_bullet(doc, "Graph Neural Networks: Model market connectivity and spatial price transmission using GNNs on the Maharashtra market network graph.")
    add_bullet(doc, "Hierarchical Reconciliation: Ensure forecasts at market level are consistent with district and state-level aggregates.")
    add_bullet(doc, "Crop Disease Identification: Integrate image-based crop disease detection using phone camera to complement the veterinary services module.")
    add_bullet(doc, "Government Scheme Integration: Display relevant agricultural schemes (PM-KISAN, crop insurance) based on farmer profile and location.")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# CHAPTER 9 - CONCLUSION
# ══════════════════════════════════════════════════════════════
def chapter9(doc):
    add_heading_styled(doc, "Chapter 9: Conclusion", 1)

    body_para(doc,
        "MANDIMITRA successfully demonstrates that a comprehensive, Maharashtra-focused "
        "agricultural intelligence platform can be built by combining modern data engineering "
        "practices with machine learning and full-stack web development. The project achieves "
        "all six stated objectives, as summarized below.")

    add_heading_styled(doc, "9.1 Objective Achievement", 2)
    add_bullet(doc, "Data Pipeline (Objective 1): A production-quality pipeline processes 6.1 million+ mandi records from AGMARKNET and Kaggle with DuckDB-based deduplication, removing ~700K duplicates using a 7-column natural key. Pandera schema validation and strict Maharashtra-only filtering (using filters[state.keyword]) ensure data integrity. The pipeline is resumable, parallel-capable (ThreadPoolExecutor), and generates comprehensive markdown audit reports.")
    add_bullet(doc, "Weather Integration (Objective 2): Weather data from NASA POWER (10 years of daily historical data: precipitation, temperature, humidity) and Open-Meteo (16-day forecasts) is successfully integrated for all 36 Maharashtra district headquarters, creating ML-ready joined datasets on the (date, district) key.")
    add_bullet(doc, "Crop Risk Advisory (Objective 3): The LightGBM classifier with physics-informed features (GDD, VPD, Drought Stress Index) achieves 87.4% overall accuracy. High Risk recall was improved by 3.1% (from 42.98% to 45.89%) through focal-loss inspired class weighting (Low:1, Medium:10, High:50). The GDD feature achieved rank 9 in feature importance, validating the physics-informed approach.")
    add_bullet(doc, "Price Forecasting (Objective 4): The 5-horizon Price Intelligence Engine achieves R-squared = 0.93 for 1-day predictions and maintains above 0.86 at all horizons. Residual-based conformal prediction intervals provide calibrated uncertainty at 80%, 90%, and 95% confidence levels, enabling conservative HOLD/SELL recommendations.")
    add_bullet(doc, "Web Application (Objective 5): A full-stack application built with Next.js 14, FastAPI, and Supabase provides six main pages and three role-based dashboards with modern UI/UX including Framer Motion animations, Tailwind CSS responsive design, and real-time data updates. Authentication uses Supabase Auth with JWT Bearer tokens.")
    add_bullet(doc, "Veterinary Services (Objective 6): A complete veterinary module with 15 API endpoints enables doctor verification (document upload + admin review), appointment booking (with time slot selection from 8 slots), booking status management (confirm/complete/cancel), and emergency SOS broadcasts with first-come-first-serve doctor acceptance.")

    add_heading_styled(doc, "9.2 Technical Contributions", 2)
    body_para(doc,
        "The key technical contributions of MANDIMITRA that advance the state of practice "
        "in agricultural technology are:")
    add_bullet(doc, "Physics-Informed ML for Agriculture: The integration of Growing Degree Days, Vapor Pressure Deficit, and Drought Stress Index into a tabular ML model, validated by feature importance analysis, demonstrates the value of domain knowledge in agricultural risk prediction.")
    add_bullet(doc, "Conformal Prediction for Farmers: Applying residual-based conformal prediction to agricultural price forecasting provides farmers with calibrated uncertainty estimates - a feature absent in all existing agricultural platforms.")
    add_bullet(doc, "Unified Agri-Vet Platform: MANDIMITRA is the first platform to combine agricultural market intelligence (prices, risk advisory, forecasting) with veterinary services (verification, booking, emergency SOS) in a single application designed for Maharashtra farmers.")
    add_bullet(doc, "Data Pipeline Engineering: The production-quality pipeline with adaptive rate limiting, chunked resumable downloads, DuckDB-based deduplication, and comprehensive audit reporting sets a standard for agricultural data engineering.")

    add_heading_styled(doc, "9.3 Deployment Architecture", 2)
    add_image(doc, "deployment_architecture.png", "Figure 9.1: Production Deployment Architecture")
    body_para(doc,
        "MANDIMITRA is designed for production deployment with the architecture shown in "
        "Figure 9.1. The frontend can be deployed to Vercel (free tier) as a static export. "
        "The FastAPI backend runs on a Python server (VPS or cloud instance). Supabase "
        "handles authentication, database, and file storage as a managed service. ML model "
        "files (.joblib) are loaded at server startup for fast inference.")

    add_heading_styled(doc, "9.4 Final Remarks", 2)
    body_para(doc,
        "MANDIMITRA demonstrates that it is feasible to build a comprehensive agricultural "
        "intelligence platform using entirely open-source tools and free-tier cloud services, "
        "making it financially accessible for deployment in developing regions. The modular "
        "architecture ensures independent scaling and maintenance of each component. With "
        "the identified future enhancements (mobile app, Marathi language, advanced ML "
        "models), MANDIMITRA has the potential to serve as a model for state-specific "
        "agricultural intelligence systems across India, ultimately improving the "
        "livelihoods of millions of farmers.")
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
def references_page(doc):
    add_heading_styled(doc, "References", 1)

    refs = [
        'Data.gov.in - AGMARKNET Agricultural Marketing, "Daily Market Prices," https://data.gov.in, Government of India, 2001-2026.',
        'NASA POWER Project, "Prediction of Worldwide Energy Resources," https://power.larc.nasa.gov/, NASA Langley Research Center.',
        'Open-Meteo, "Free Weather Forecast API," https://open-meteo.com/.',
        'Kaggle, "India Agricultural Commodity Prices (AGMARKNET Archive)," https://www.kaggle.com/datasets/, 2001-2024.',
        'G. Ke, Q. Meng, T. Finley et al., "LightGBM: A Highly Efficient Gradient Boosting Decision Tree," NeurIPS, 2017.',
        'S. Patel et al., "MT-CYP-Net: Multi-Task Network for Crop Yield Prediction using Remote Sensing Data," arXiv:2505.12069, 2025.',
        'R. Kumar et al., "NeuralCrop: Combining Physics and Machine Learning for Crop Yield Predictions," arXiv:2512.20177, 2025.',
        'A. Sharma et al., "Intrinsic Explainability of Multimodal Learning for Crop Yield Prediction," arXiv:2508.06939, 2025.',
        'N. Hollmann et al., "TabPFN: A Transformer That Solves Small Tabular Classification Problems in a Second," arXiv:2207.01848, NeurIPS 2022.',
        'M. Weber et al., "Explainability of Sub-Field Level Crop Yield Prediction using Machine Learning," arXiv:2407.08274, ICML 2024.',
        'Y. Romano, E. Patterson, R.J. Tibshirani, "Conformalized Quantile Regression," NeurIPS, 2019.',
        'Supabase, "The Open Source Firebase Alternative," https://supabase.com/.',
        'Next.js, "The React Framework for Production," https://nextjs.org/, Vercel Inc.',
        'FastAPI, "Modern, Fast Web Framework for Building APIs with Python 3.7+," https://fastapi.tiangolo.com/, S. Ramirez.',
        'Tailwind CSS, "A Utility-First CSS Framework for Rapid UI Development," https://tailwindcss.com/.',
        'DuckDB, "An In-Process SQL OLAP Database Management System," https://duckdb.org/.',
        'Framer Motion, "A Production-Ready Motion Library for React," https://www.framer.com/motion/.',
        'Maharashtra State Agricultural Marketing Board (MSAMB), "Market Information," https://www.msamb.com/.',
        'Ministry of Agriculture and Farmers Welfare, Government of India, "e-NAM: National Agriculture Market," https://www.enam.gov.in/.',
        '20th Livestock Census, "All India Report," Department of Animal Husbandry and Dairying, Government of India, 2019.',
    ]
    for i, r in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        run = p.add_run(f"[{i}] {r}")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# APPENDICES
# ══════════════════════════════════════════════════════════════
def appendices_page(doc):
    add_heading_styled(doc, "Appendices", 1)

    add_heading_styled(doc, "Appendix A: Project Directory Structure", 2)
    add_code_block(doc,
        "mandimitra/\n"
        "  configs/\n"
        "    project.yaml\n"
        "    data_sources.yaml\n"
        "    maharashtra_locations.csv (36 districts)\n"
        "    crop_lifecycle.json\n"
        "  data/\n"
        "    raw/mandi/   (chunked by district)\n"
        "    raw/weather/ (power_daily, openmeteo)\n"
        "    processed/mandi/  (canonical, merged)\n"
        "    processed/weather/(power, forecast)\n"
        "    processed/model/  (ML-ready datasets)\n"
        "    metadata/maharashtra/ (discovery)\n"
        "  scripts/\n"
        "    download_all_data.py\n"
        "    download_mandi_history_kaggle.py\n"
        "    download_mandi_current_datagov.py\n"
        "    merge_mandi_datasets.py\n"
        "    download_weather_power_maharashtra.py\n"
        "    download_weather_openmeteo_maharashtra.py\n"
        "    build_all_processed.py\n"
        "    build_canonical_mandi.py\n"
        "    process_weather.py\n"
        "    build_model_datasets.py\n"
        "    train_crop_risk_model.py\n"
        "    train_price_model.py\n"
        "    validate_data.py\n"
        "    self_check.py\n"
        "  src/\n"
        "    schemas/ (mandi.py, weather.py)\n"
        "    utils/   (http, io, logging, maharashtra)\n"
        "  api/\n"
        "    main.py  (FastAPI app entry)\n"
        "    auth.py  (signup, login, profiles)\n"
        "    vet.py   (15 vet endpoints)\n"
        "  web/src/app/\n"
        "    page.tsx         (landing page)\n"
        "    crop-risk/       (advisory page)\n"
        "    price-forecast/  (prediction page)\n"
        "    veterinary/      (vet services info)\n"
        "    login/ signup/   (auth pages)\n"
        "    dashboard/\n"
        "      farmer/page.tsx  (farmer dashboard)\n"
        "      doctor/page.tsx  (doctor dashboard)\n"
        "      admin/page.tsx   (admin dashboard)\n"
        "  web/src/components/sections/\n"
        "    HeroSection.tsx\n"
        "    FeaturesSection.tsx\n"
        "    StatsSection.tsx\n"
        "    HowItWorksSection.tsx\n"
        "    TestimonialsSection.tsx\n"
        "    CTASection.tsx\n"
        "  web/src/contexts/\n"
        "    AuthContext.tsx\n"
        "  models/\n"
        "    crop_risk_advisor/  (model.joblib, metrics)\n"
        "    price_intelligence/ (5 horizon models)\n"
        "  reports/\n"
        "    diagrams/ (18 PNG figures at 300 DPI)\n"
        "    generate_diagrams.py\n"
        "    generate_report.py\n"
        "    generate_report_docx.py")

    add_heading_styled(doc, "Appendix B: Maharashtra Districts (36)", 2)
    districts = [
        ["1", "Ahmednagar", "West"], ["2", "Akola", "Vidarbha"], ["3", "Amravati", "Vidarbha"],
        ["4", "Aurangabad", "Marathwada"], ["5", "Beed", "Marathwada"], ["6", "Bhandara", "Vidarbha"],
        ["7", "Buldhana", "Vidarbha"], ["8", "Chandrapur", "Vidarbha"], ["9", "Dhule", "North"],
        ["10", "Gadchiroli", "Vidarbha"], ["11", "Gondia", "Vidarbha"], ["12", "Hingoli", "Marathwada"],
        ["13", "Jalgaon", "North"], ["14", "Jalna", "Marathwada"], ["15", "Kolhapur", "West"],
        ["16", "Latur", "Marathwada"], ["17", "Mumbai", "Konkan"], ["18", "Mumbai Sub.", "Konkan"],
        ["19", "Nagpur", "Vidarbha"], ["20", "Nanded", "Marathwada"], ["21", "Nandurbar", "North"],
        ["22", "Nashik", "North"], ["23", "Osmanabad", "Marathwada"], ["24", "Palghar", "Konkan"],
        ["25", "Parbhani", "Marathwada"], ["26", "Pune", "West"], ["27", "Raigad", "Konkan"],
        ["28", "Ratnagiri", "Konkan"], ["29", "Sangli", "West"], ["30", "Satara", "West"],
        ["31", "Sindhudurg", "Konkan"], ["32", "Solapur", "West"], ["33", "Thane", "Konkan"],
        ["34", "Wardha", "Vidarbha"], ["35", "Washim", "Vidarbha"], ["36", "Yavatmal", "Vidarbha"],
    ]
    add_table(doc, ["#", "District", "Region"], districts, caption="Table B.1: All 36 Maharashtra Districts with Region")

    add_heading_styled(doc, "Appendix C: API Endpoint Quick Reference", 2)
    body_para(doc, "All endpoints are prefixed with /api/vet and require Bearer token.", size=11)
    add_code_block(doc,
        "# ADMIN (3 endpoints)\n"
        "GET  /admin/pending-doctors     -> {doctors: [...]}\n"
        "POST /admin/verify-doctor       <- {doctor_id, action}\n"
        "GET  /admin/stats               -> {total_farmers, ...}\n\n"
        "# DOCTOR (7 endpoints)\n"
        "POST /doctor/upload-document    <- FormData(file)\n"
        "GET  /doctor/profile            -> {profile, stats}\n"
        "GET  /doctor/emergency-cases    -> {emergencies: [...]}\n"
        "POST /doctor/accept-emergency   <- {emergency_id}\n"
        "POST /doctor/complete-emergency <- {emergency_id}\n"
        "GET  /doctor/bookings           -> {bookings: [...]}\n"
        "PATCH /doctor/booking-status    <- {booking_id, status}\n\n"
        "# FARMER (5 endpoints)\n"
        "GET  /doctors                   -> {doctors: [...]}\n"
        "POST /farmer/book              <- {doctor_id, date, slot}\n"
        "POST /farmer/emergency          <- {animal, desc, loc}\n"
        "GET  /farmer/bookings           -> {bookings: [...]}\n"
        "GET  /farmer/emergencies        -> {emergencies: [...]}")

    add_heading_styled(doc, "Appendix D: Key Configuration Files", 2)
    add_heading_styled(doc, "D.1 project.yaml (excerpt)", 3)
    add_code_block(doc,
        "project:\n"
        "  name: mandimitra\n"
        "  version: 1.0.0\n\n"
        "maharashtra:\n"
        "  state_name: Maharashtra\n"
        "  state_code: MH\n"
        "  total_districts: 36\n\n"
        "mandi:\n"
        "  resource_id: 9ef84268-d588-465a-...\n"
        "  page_size: 1000\n"
        "  state_filter: Maharashtra  # LOCKED\n"
        "  chunk_by: district\n\n"
        "nasa_power:\n"
        "  parameters: [PRECTOTCORR, T2M, RH2M]\n"
        "  default_days_back: 365\n\n"
        "openmeteo:\n"
        "  forecast_days: 16\n"
        "  timezone: Asia/Kolkata")


# ══════════════════════════════════════════════════════════════
# BUILD
# ══════════════════════════════════════════════════════════════
def main():
    print("Building MANDIMITRA MSBTE Report (.docx Word Document)...")
    print(f"  Output: {OUTPUT}")
    print(f"  Logo: {LOGO} (exists={LOGO.exists()})")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Build document
    cover_page(doc)
    certificate_page(doc)
    acknowledgement_page(doc)
    abstract_page(doc)
    list_of_figures(doc)
    list_of_tables(doc)
    table_of_contents(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    chapter7(doc)
    chapter8(doc)
    chapter9(doc)
    references_page(doc)
    appendices_page(doc)

    # Add footer to the document
    add_footer(doc)

    doc.save(str(OUTPUT))

    size = os.path.getsize(OUTPUT)
    print(f"\n  Generated: {size:,} bytes")
    print("  Done.")


if __name__ == "__main__":
    main()
