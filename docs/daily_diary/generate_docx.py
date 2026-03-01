"""Generate .docx Daily Diary files for all 4 students."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Student Data ──────────────────────────────────────────────
STUDENTS = [
    {
        "name": "Avishkar Borate",
        "enrolment": "01",
        "role": "Data Pipeline & Backend Development",
        "weeks": [
            {  # Week 1
                "week": 1,
                "rows": [
                    ("Monday",    "08/12/25", "Review of problem statement and project objective", "Clear understanding of MANDIMITRA scope and goals"),
                    ("Tuesday",   "09/12/25", "Study of AGMARKNET API documentation and data.gov.in portal", "API structure and data fields understood"),
                    ("Wednesday", "10/12/25", "Literature survey on data pipeline architectures for agricultural data", "Knowledge of ETL patterns for large-scale data gained"),
                    ("Thursday",  "11/12/25", "Analysis of existing agricultural data platforms and their limitations", "Performance gaps and feature gaps identified"),
                    ("Friday",    "12/12/25", "Environment setup — Python 3.10, virtual environment, project structure", "Dev environment configured with all dependencies"),
                ],
            },
            {  # Week 2
                "week": 2,
                "rows": [
                    ("Monday",    "15/12/25", "Study of Data.gov.in REST API authentication and rate limiting", "API key registration and request patterns understood"),
                    ("Tuesday",   "16/12/25", "Initial prototype for AGMARKNET data download script", "Basic API call to fetch mandi prices working"),
                    ("Wednesday", "17/12/25", "Research on Maharashtra districts, mandis, and commodity classification", "36 district mapping and commodity list finalized"),
                    ("Thursday",  "18/12/25", "Design of project.yaml and data_sources.yaml configuration schema", "Config-driven pipeline architecture designed"),
                    ("Friday",    "19/12/25", "Setup of .env template and secure API key management", "API key handling with no hardcoded secrets achieved"),
                ],
            },
            {  # Week 3
                "week": 3,
                "rows": [
                    ("Monday",    "22/12/25", "Implementation of chunked download system for AGMARKNET data", "Resumable by-district chunking logic implemented"),
                    ("Tuesday",   "23/12/25", "Development of adaptive rate limiting with token bucket algorithm", "429 error handling and backoff strategy working"),
                    ("Wednesday", "24/12/25", "Building Maharashtra-only validation filters — reject non-MH data", "Strict state constraint enforced in pipeline"),
                    ("Thursday",  "25/12/25", "Implementation of district name normalization (35 raw → 36 canonical)", "District mapping table with fuzzy matching complete"),
                    ("Friday",    "26/12/25", "Creation of Pandera validation schemas for mandi price data", "Schema validation with strict type and range checks done"),
                ],
            },
            {  # Week 4
                "week": 4,
                "rows": [
                    ("Monday",    "29/12/25", "Integration of NASA POWER API for historical rainfall data", "Daily precipitation for 36 district HQs downloading"),
                    ("Tuesday",   "30/12/25", "Development of Open-Meteo weather forecast integration", "16-day rainfall forecast for Maharashtra districts working"),
                    ("Wednesday", "31/12/25", "Building parallel download with ThreadPoolExecutor", "Configurable worker pool for concurrent API calls done"),
                    ("Thursday",  "01/01/26", "Implementation of atomic progress tracking for downloads", "Resume-safe download state using JSON checkpoints"),
                    ("Friday",    "02/01/26", "Creation of markdown audit reports for data compliance", "Automated audit log generation for every pipeline run"),
                ],
            },
            {  # Week 5
                "week": 5,
                "rows": [
                    ("Monday",    "05/01/26", "Design of FastAPI backend architecture and route structure", "API module structure with auth, crop, weather, vet routes planned"),
                    ("Tuesday",   "06/01/26", "Implementation of FastAPI main application with CORS and middleware", "api/main.py with health check and CORS configured"),
                    ("Wednesday", "07/01/26", "Development of authentication module using Supabase", "api/auth.py with JWT-based user auth implemented"),
                    ("Thursday",  "08/01/26", "Building weather and market data API endpoints", "api/weather_market.py serving real-time price and weather data"),
                    ("Friday",    "09/01/26", "Implementation of veterinary emergency services API", "api/vet.py with nearby vet lookup and emergency contacts"),
                ],
            },
            {  # Week 6
                "week": 6,
                "rows": [
                    ("Monday",    "12/01/26", "Supabase PostgreSQL database schema design and table creation", "supabase_setup.sql with all required tables written"),
                    ("Tuesday",   "13/01/26", "Implementation of data deduplication logic in pipeline", "Dedup flowchart implemented — no duplicate records in DB"),
                    ("Wednesday", "14/01/26", "Development of crop disease detection API endpoint", "api/crop_disease.py with image upload and prediction serving"),
                    ("Thursday",  "15/01/26", "Building Kaggle dataset integration for historical mandi data", "Kaggle API integration for bulk historical data download"),
                    ("Friday",    "16/01/26", "Data merge pipeline — combining current and historical datasets", "Merge logic with conflict resolution and timestamp handling"),
                ],
            },
            {  # Week 7
                "week": 7,
                "rows": [
                    ("Monday",    "19/01/26", "Optimization of data pipeline for 3.5M+ records processing", "Memory-safe streaming with constant memory usage achieved"),
                    ("Tuesday",   "20/01/26", "Implementation of discovery mode for new mandis and commodities", "Streaming discovery without loading full dataset in memory"),
                    ("Wednesday", "21/01/26", "Building data quality validation and completeness checks", "Automated data completeness reports generated"),
                    ("Thursday",  "22/01/26", "Supabase v2 and v3 schema migrations for enhanced features", "supabase_setup_v2.sql and v3.sql with new columns and indexes"),
                    ("Friday",    "23/01/26", "API endpoint optimization — response caching and query tuning", "Reduced API response time by 40% with query optimization"),
                ],
            },
            {  # Week 8
                "week": 8,
                "rows": [
                    ("Monday",    "26/01/26", "Integration testing of full data pipeline end-to-end", "All pipeline stages passing with validated output"),
                    ("Tuesday",   "27/01/26", "Backend API integration with ML model serving endpoints", "FastAPI serving predictions from all 3 trained models"),
                    ("Wednesday", "28/01/26", "Error handling and logging improvements across backend", "Structured logging with log files in logs/ directory"),
                    ("Thursday",  "29/01/26", "Implementation of locations.csv and maharashtra_locations.csv", "District HQ coordinates for 36 districts configured"),
                    ("Friday",    "30/01/26", "Load testing of FastAPI backend with concurrent requests", "Backend handling 100+ concurrent requests without errors"),
                ],
            },
            {  # Week 9
                "week": 9,
                "rows": [
                    ("Monday",    "02/02/26", "Dockerfile creation for containerized backend deployment", "Multi-stage Docker build with optimized image size"),
                    ("Tuesday",   "03/02/26", "Render.yaml configuration for cloud deployment", "render.yaml with environment variables and build commands"),
                    ("Wednesday", "04/02/26", "Requirements file optimization — separate server vs full deps", "requirements-server.txt with minimal production dependencies"),
                    ("Thursday",  "05/02/26", "Data pipeline automation — scheduled downloads and validation", "Automated nightly data refresh pipeline configured"),
                    ("Friday",    "06/02/26", "API security hardening — input validation, rate limiting", "All endpoints sanitized with proper error responses"),
                ],
            },
            {  # Week 10
                "week": 10,
                "rows": [
                    ("Monday",    "09/02/26", "Backend deployment to Render cloud platform", "Live API serving at production URL on Render"),
                    ("Tuesday",   "10/02/26", "Database migration to production Supabase instance", "Production database with all tables and seed data ready"),
                    ("Wednesday", "11/02/26", "Monitoring and health check endpoint implementation", "/health and /status endpoints for uptime monitoring"),
                    ("Thursday",  "12/02/26", "Bug fixes for edge cases in data pipeline — missing values, encoding", "All known data quality issues resolved"),
                    ("Friday",    "13/02/26", "Performance benchmarking of production backend", "API latency <200ms for all endpoints documented"),
                ],
            },
            {  # Week 11
                "week": 11,
                "rows": [
                    ("Monday",    "16/02/26", "Code review and refactoring of pipeline and API modules", "Clean, documented codebase with consistent style"),
                    ("Tuesday",   "17/02/26", "Writing inline documentation and API docstrings", "All functions and endpoints documented with examples"),
                    ("Wednesday", "18/02/26", "Final data pipeline run with full Maharashtra dataset", "Complete 3.5M+ record dataset validated and stored"),
                    ("Thursday",  "19/02/26", "Backend stress testing and edge case verification", "All edge cases handled — empty data, invalid inputs, timeouts"),
                    ("Friday",    "20/02/26", "Preparation of backend architecture diagrams for report", "DFD Level-0, pipeline flowchart, module interaction diagrams"),
                ],
            },
            {  # Week 12
                "week": 12,
                "rows": [
                    ("Monday",    "23/02/26", "Final integration testing with frontend and ML models", "Full stack working end-to-end in production"),
                    ("Tuesday",   "24/02/26", "README.md finalization with complete setup and usage instructions", "Comprehensive README with quick start, config, and API docs"),
                    ("Wednesday", "25/02/26", "Demo preparation — live data pipeline walkthrough", "Demo script for capstone presentation ready"),
                    ("Thursday",  "26/02/26", "Practice presentation of backend and pipeline components", "Confident delivery of technical explanation achieved"),
                    ("Friday",    "27/02/26", "Final submission — code freeze, backup, and handover", "All backend deliverables submitted and archived"),
                ],
            },
        ],
    },
    {
        "name": "Amol Salunke",
        "enrolment": "02",
        "role": "Machine Learning & Model Training",
        "weeks": [
            {
                "week": 1,
                "rows": [
                    ("Monday",    "08/12/25", "Review of problem statement and project objective", "Clear understanding of MANDIMITRA AI requirements"),
                    ("Tuesday",   "09/12/25", "Study of previously implemented ML models in agriculture", "Individual model roles and architectures understood"),
                    ("Wednesday", "10/12/25", "Literature survey on ensemble and hybrid models (MT-CYP-Net, NeuralCrop)", "Knowledge of hybrid ML techniques for agriculture gained"),
                    ("Thursday",  "11/12/25", "Analysis of model performance metrics — precision, recall, F1, AUC", "Performance evaluation criteria for all models identified"),
                    ("Friday",    "12/12/25", "Environment setup for ML — TensorFlow, LightGBM, scikit-learn, pandas", "Required ML libraries installed and configured"),
                ],
            },
            {
                "week": 2,
                "rows": [
                    ("Monday",    "15/12/25", "Exploratory data analysis of mandi price dataset (3.5M+ records)", "Data distributions, seasonal patterns, and outliers understood"),
                    ("Tuesday",   "16/12/25", "Study of LightGBM algorithm for tabular data prediction", "LightGBM hyperparameters and training pipeline understood"),
                    ("Wednesday", "17/12/25", "Research on MobileNetV2 architecture for image classification", "Transfer learning approach for disease detection planned"),
                    ("Thursday",  "18/12/25", "Analysis of crop lifecycle data from crop_lifecycle.json", "Crop stages, growth periods, and risk windows documented"),
                    ("Friday",    "19/12/25", "Data preprocessing plan — missing values, encoding, normalization", "Complete preprocessing pipeline for all models designed"),
                ],
            },
            {
                "week": 3,
                "rows": [
                    ("Monday",    "22/12/25", "Feature engineering — physics-informed features (GDD, VPD, Drought Index)", "Growing Degree Days and Vapour Pressure Deficit features computed"),
                    ("Tuesday",   "23/12/25", "Feature engineering — temporal features (seasonality, lag, rolling stats)", "30-day rolling averages and seasonal decomposition added"),
                    ("Wednesday", "24/12/25", "Data splitting strategy — train/validation/test with temporal ordering", "Time-aware split ensuring no data leakage implemented"),
                    ("Thursday",  "25/12/25", "Initial Crop Risk Advisor model training with LightGBM", "Baseline model achieving 72% accuracy on validation set"),
                    ("Friday",    "26/12/25", "Focal-loss inspired class weighting for high-risk crop detection", "Class weighting scheme from MT-CYP-Net paper applied"),
                ],
            },
            {
                "week": 4,
                "rows": [
                    ("Monday",    "29/12/25", "Hyperparameter tuning of Crop Risk Advisor — learning rate, depth, leaves", "Optimized LightGBM config with best validation score"),
                    ("Tuesday",   "30/12/25", "Crop Risk Advisor evaluation — confusion matrix, classification report", "High-risk recall improved by 3.1% with focal-loss weighting"),
                    ("Wednesday", "31/12/25", "Design of Price Intelligence Engine — multi-horizon prediction approach", "7-day, 14-day, and 30-day forecast architecture planned"),
                    ("Thursday",  "01/01/26", "Feature engineering for price prediction — price lags, volume, seasonal", "Commodity-specific feature sets for price models created"),
                    ("Friday",    "02/01/26", "Initial Price Intelligence Engine training for Maharashtra commodities", "Baseline price prediction model trained on top 10 commodities"),
                ],
            },
            {
                "week": 5,
                "rows": [
                    ("Monday",    "05/01/26", "Conformal prediction implementation for prediction intervals", "80/90/95% confidence intervals on price forecasts working"),
                    ("Tuesday",   "06/01/26", "Multi-horizon price model training (7-day, 14-day, 30-day)", "Three separate forecast models trained and validated"),
                    ("Wednesday", "07/01/26", "Price model evaluation — MAE, RMSE, and interval coverage", "Coverage rates: 80% interval = 82.3%, 90% = 91.1%, 95% = 96.2%"),
                    ("Thursday",  "08/01/26", "Preparation of Crop Disease dataset — 17 classes, image organization", "17-class dataset with Corn, Potato, Rice, Sugarcane, Wheat organized"),
                    ("Friday",    "09/01/26", "Image augmentation pipeline — rotation, flip, zoom, brightness", "Training data augmented 5x with realistic transformations"),
                ],
            },
            {
                "week": 6,
                "rows": [
                    ("Monday",    "12/01/26", "MobileNetV2 transfer learning — loading pretrained ImageNet weights", "Base model loaded with frozen feature extraction layers"),
                    ("Tuesday",   "13/01/26", "Custom classification head design — GlobalAveragePooling + Dense layers", "17-class output head with dropout regularization added"),
                    ("Wednesday", "14/01/26", "Disease detection model training — 50 epochs with early stopping", "Training accuracy 94.2%, validation accuracy 91.8% achieved"),
                    ("Thursday",  "15/01/26", "Fine-tuning top layers of MobileNetV2 for domain adaptation", "Validation accuracy improved to 93.5% after fine-tuning"),
                    ("Friday",    "16/01/26", "Model evaluation — per-class precision, recall, and confusion matrix", "All 17 classes achieving >85% F1 score confirmed"),
                ],
            },
            {
                "week": 7,
                "rows": [
                    ("Monday",    "19/01/26", "Model size optimization — quantization and pruning of disease detector", "Model size reduced from 22 MB to ~14 MB with minimal accuracy loss"),
                    ("Tuesday",   "20/01/26", "TensorFlow Lite conversion for mobile-optimized inference", "TFLite model generated for on-device prediction"),
                    ("Wednesday", "21/01/26", "TFJS conversion for web-based disease detection", "TensorFlow.js model exported for browser inference"),
                    ("Thursday",  "22/01/26", "Model serving pipeline — FastAPI endpoint for image prediction", "/predict endpoint accepting image upload and returning disease class"),
                    ("Friday",    "23/01/26", "Batch inference testing on unseen disease images", "Model correctly identifying diseases on 50 new test images"),
                ],
            },
            {
                "week": 8,
                "rows": [
                    ("Monday",    "26/01/26", "Crop Risk Advisor integration with weather forecast features", "Risk model using live 16-day weather forecast for prediction"),
                    ("Tuesday",   "27/01/26", "Price Intelligence Engine integration with latest mandi data", "Price model predictions updated with most recent market data"),
                    ("Wednesday", "28/01/26", "Cross-validation of all 3 models with k-fold strategy", "Consistent performance across folds — no overfitting detected"),
                    ("Thursday",  "29/01/26", "Model versioning and artifact management in models/ directory", "All model files, configs, and metrics saved with version tags"),
                    ("Friday",    "30/01/26", "End-to-end ML pipeline testing — data → features → prediction", "Complete ML pipeline running without errors on fresh data"),
                ],
            },
            {
                "week": 9,
                "rows": [
                    ("Monday",    "02/02/26", "Model optimization plan documentation", "MODEL_OPTIMIZATION_PLAN.md with all strategies documented"),
                    ("Tuesday",   "03/02/26", "Optimization execution — feature selection, hyperparameter refinement", "Top features identified, redundant features removed"),
                    ("Wednesday", "04/02/26", "Optimization results documentation and benchmarking", "OPTIMIZATION_RESULTS.md with before/after comparisons"),
                    ("Thursday",  "05/02/26", "Crop risk training report generation with detailed metrics", "crop_risk_training_report.md with all training details"),
                    ("Friday",    "06/02/26", "Model robustness testing — adversarial inputs, edge cases", "Models handling noisy data, missing values gracefully"),
                ],
            },
            {
                "week": 10,
                "rows": [
                    ("Monday",    "09/02/26", "Production model deployment — model files to Render server", "All 3 models serving predictions in production"),
                    ("Tuesday",   "10/02/26", "Model inference latency optimization in production", "Disease detection <500ms, price/risk prediction <100ms"),
                    ("Wednesday", "11/02/26", "A/B testing of model versions with production traffic", "Latest model versions outperforming baselines confirmed"),
                    ("Thursday",  "12/02/26", "Model monitoring — tracking prediction distribution shifts", "Alerting setup for data drift detection"),
                    ("Friday",    "13/02/26", "Model documentation — architecture, training, evaluation details", "Complete model cards for all 3 models written"),
                ],
            },
            {
                "week": 11,
                "rows": [
                    ("Monday",    "16/02/26", "User feedback analysis on model predictions", "Prediction accuracy validated by sample farmer feedback"),
                    ("Tuesday",   "17/02/26", "Retraining models with additional data collected in January", "Model accuracy improved with larger training dataset"),
                    ("Wednesday", "18/02/26", "Final model evaluation and comparison table for report", "Chapter 6 results table with all metrics prepared"),
                    ("Thursday",  "19/02/26", "ML training flow and results diagrams for capstone report", "ml_training_flow.png, crop_risk_results.png created"),
                    ("Friday",    "20/02/26", "Preparation of ML pipeline walkthrough for capstone demo", "Live demo of model training and prediction ready"),
                ],
            },
            {
                "week": 12,
                "rows": [
                    ("Monday",    "23/02/26", "Final model accuracy verification on production data", "All models meeting target accuracy thresholds"),
                    ("Tuesday",   "24/02/26", "Model export and backup — SavedModel, TFLite, TFJS formats", "All model formats archived for submission"),
                    ("Wednesday", "25/02/26", "Literature survey comparison table for capstone report", "lit_survey_comparison.png with paper vs implementation mapping"),
                    ("Thursday",  "26/02/26", "Practice presentation of ML components and results", "Confident explanation of model choices and performance"),
                    ("Friday",    "27/02/26", "Final submission — model artifacts, training logs, evaluation reports", "All ML deliverables submitted and archived"),
                ],
            },
        ],
    },
    {
        "name": "Pruthviraj Hippirkar",
        "enrolment": "03",
        "role": "Frontend & Mobile App Development",
        "weeks": [
            {
                "week": 1,
                "rows": [
                    ("Monday",    "08/12/25", "Review of problem statement and project objective", "Clear understanding of user-facing features required"),
                    ("Tuesday",   "09/12/25", "Study of existing agricultural apps and farmer-facing platforms", "UX patterns and usability gaps in current apps identified"),
                    ("Wednesday", "10/12/25", "Literature survey on mobile-first design for rural users", "Low-literacy UI principles and accessibility standards studied"),
                    ("Thursday",  "11/12/25", "Analysis of MANDIMITRA feature requirements for frontend", "Dashboard, price view, disease scan, vet services scoped"),
                    ("Friday",    "12/12/25", "Environment setup — Node.js, React, Flutter SDK, VS Code extensions", "Frontend and mobile dev environments configured"),
                ],
            },
            {
                "week": 2,
                "rows": [
                    ("Monday",    "15/12/25", "Wireframe design for web dashboard — homepage, price tracker, risk view", "Low-fidelity wireframes for all 5 main screens created"),
                    ("Tuesday",   "16/12/25", "Wireframe design for mobile app — navigation, camera scan, alerts", "Mobile wireframes with bottom nav and gesture patterns done"),
                    ("Wednesday", "17/12/25", "Color scheme, typography, and design system selection", "Agriculture-themed green/earth-tone design system finalized"),
                    ("Thursday",  "18/12/25", "Study of React component architecture and state management", "Component hierarchy and data flow patterns planned"),
                    ("Friday",    "19/12/25", "Project scaffold — React app with routing and folder structure", "web/ directory initialized with pages, components, hooks"),
                ],
            },
            {
                "week": 3,
                "rows": [
                    ("Monday",    "22/12/25", "Implementation of responsive navigation bar and sidebar", "Navigation component with mobile hamburger menu working"),
                    ("Tuesday",   "23/12/25", "Homepage development — hero section, feature cards, CTA buttons", "Landing page with MANDIMITRA branding and feature overview"),
                    ("Wednesday", "24/12/25", "Dashboard layout — grid system for price cards and chart areas", "Responsive grid layout adapting to mobile/tablet/desktop"),
                    ("Thursday",  "25/12/25", "Implementation of Supabase authentication UI — login/signup forms", "User registration and login forms with validation working"),
                    ("Friday",    "26/12/25", "Auth state management — protected routes and session handling", "Authenticated routing with Supabase JWT integration done"),
                ],
            },
            {
                "week": 4,
                "rows": [
                    ("Monday",    "29/12/25", "Mandi Price Tracker page — commodity selector, district filter", "Price search UI with dropdown filters for district and crop"),
                    ("Tuesday",   "30/12/25", "Price chart component using Chart.js — line and bar charts", "Interactive price trend charts with 7/14/30 day views"),
                    ("Wednesday", "31/12/25", "Price prediction display — showing confidence intervals visually", "Shaded band charts for 80/90/95% prediction intervals"),
                    ("Thursday",  "01/01/26", "Commodity comparison view — side-by-side price analysis", "Multi-commodity comparison with synchronized date axis"),
                    ("Friday",    "02/01/26", "API integration — connecting price UI to FastAPI backend", "Live price data fetching from /api/prices endpoint working"),
                ],
            },
            {
                "week": 5,
                "rows": [
                    ("Monday",    "05/01/26", "Crop Risk Assessment page — risk score display and color coding", "Risk dashboard with High/Medium/Low cards and percentages"),
                    ("Tuesday",   "06/01/26", "Weather widget — current conditions and 16-day forecast display", "Weather card with temperature, rainfall, and forecast chart"),
                    ("Wednesday", "07/01/26", "Risk factor breakdown — detailed view of contributing factors", "Expandable risk factors showing weather, price, and crop stage"),
                    ("Thursday",  "08/01/26", "District-wise risk map component using interactive Maharashtra map", "SVG map of Maharashtra with color-coded risk districts"),
                    ("Friday",    "09/01/26", "API integration — connecting risk UI to /api/crop-risk endpoint", "Live risk assessment data flowing from backend to frontend"),
                ],
            },
            {
                "week": 6,
                "rows": [
                    ("Monday",    "12/01/26", "Crop Disease Detection page — camera capture and image upload UI", "Image upload with drag-drop and camera access on mobile"),
                    ("Tuesday",   "13/01/26", "TensorFlow.js integration for client-side disease detection", "TFJS model loading and in-browser image classification working"),
                    ("Wednesday", "14/01/26", "Disease result display — class name, confidence, treatment advice", "Result card with disease info, confidence %, and remedy text"),
                    ("Thursday",  "15/01/26", "Image preprocessing pipeline in browser — resize, normalize", "Client-side image preprocessing matching model input specs"),
                    ("Friday",    "16/01/26", "Disease history log — saving past scan results for user reference", "Scan history stored in Supabase with timestamps and results"),
                ],
            },
            {
                "week": 7,
                "rows": [
                    ("Monday",    "19/01/26", "Veterinary Emergency Services page — vet finder UI", "Search interface for nearby veterinary clinics and hospitals"),
                    ("Tuesday",   "20/01/26", "Emergency contact list — one-tap call buttons for vet helplines", "Emergency contacts with click-to-call and location display"),
                    ("Wednesday", "21/01/26", "Location-based vet search using browser geolocation API", "Auto-detecting user location for nearest vet suggestions"),
                    ("Thursday",  "22/01/26", "User profile page — settings, preferences, and notification config", "Profile management with crop preferences and district selection"),
                    ("Friday",    "23/01/26", "Notification system — price alerts and risk warnings UI", "Push notification UI for price threshold and risk alerts"),
                ],
            },
            {
                "week": 8,
                "rows": [
                    ("Monday",    "26/01/26", "Flutter mobile app setup — project structure, dependencies", "Flutter project initialized with required packages"),
                    ("Tuesday",   "27/01/26", "Mobile app — bottom navigation, routing, and theme setup", "Material Design 3 theme with agriculture green palette"),
                    ("Wednesday", "28/01/26", "Mobile app — price tracker screen with API integration", "Price view screen fetching live data from backend API"),
                    ("Thursday",  "29/01/26", "Mobile app — camera-based disease detection screen", "Camera capture with TFLite model for on-device inference"),
                    ("Friday",    "30/01/26", "Mobile app — risk assessment and weather display screens", "Risk and weather screens with pull-to-refresh pattern"),
                ],
            },
            {
                "week": 9,
                "rows": [
                    ("Monday",    "02/02/26", "Responsive design testing — mobile, tablet, desktop breakpoints", "All web pages rendering correctly on all screen sizes"),
                    ("Tuesday",   "03/02/26", "Accessibility improvements — ARIA labels, keyboard navigation, contrast", "WCAG 2.1 AA compliance for all interactive elements"),
                    ("Wednesday", "04/02/26", "Loading states, error handling, and empty state UI components", "Skeleton loaders, error banners, and empty state illustrations"),
                    ("Thursday",  "05/02/26", "Performance optimization — lazy loading, code splitting, image compression", "Lighthouse performance score improved from 65 to 92"),
                    ("Friday",    "06/02/26", "Offline capability — service worker and cache strategy for key pages", "PWA support with offline access to last-viewed data"),
                ],
            },
            {
                "week": 10,
                "rows": [
                    ("Monday",    "09/02/26", "Web app deployment — build optimization and static hosting", "Production build deployed with optimized bundle size"),
                    ("Tuesday",   "10/02/26", "Mobile APK build — release signing and APK generation", "mandimitra-release.apk built and signed for distribution"),
                    ("Wednesday", "11/02/26", "Cross-browser testing — Chrome, Firefox, Safari, Edge", "All features working consistently across browsers"),
                    ("Thursday",  "12/02/26", "Mobile device testing — Android phones and tablets", "APK tested on 5+ Android devices with different screen sizes"),
                    ("Friday",    "13/02/26", "UI bug fixes from user testing feedback", "All visual and interaction bugs resolved"),
                ],
            },
            {
                "week": 11,
                "rows": [
                    ("Monday",    "16/02/26", "User experience testing with sample users (classmates as proxies)", "Usability feedback collected and priority fixes listed"),
                    ("Tuesday",   "17/02/26", "UI polish — animations, transitions, micro-interactions", "Smooth page transitions and button feedback animations added"),
                    ("Wednesday", "18/02/26", "Final API integration testing — all frontend-backend connections", "All 12 API endpoints connected and returning correct data"),
                    ("Thursday",  "19/02/26", "Screenshots and screen recordings for project report", "All app screens captured for Chapter 5 and Chapter 6"),
                    ("Friday",    "20/02/26", "Architecture and use case diagrams for frontend in report", "use_cases_diagram.png and module_interaction.png updated"),
                ],
            },
            {
                "week": 12,
                "rows": [
                    ("Monday",    "23/02/26", "Final UI review and consistency check across all pages", "Consistent styling, spacing, and typography verified"),
                    ("Tuesday",   "24/02/26", "Demo preparation — live app walkthrough on web and mobile", "Demo flow covering all 5 major features scripted"),
                    ("Wednesday", "25/02/26", "Capstone PPT slides preparation for frontend/mobile sections", "Slide 8 (Implementation) and Slide 10 (Demo) content created"),
                    ("Thursday",  "26/02/26", "Practice presentation of frontend and mobile app components", "Confident live demo of web dashboard and mobile app"),
                    ("Friday",    "27/02/26", "Final submission — web build, APK, source code, and screenshots", "All frontend deliverables submitted and archived"),
                ],
            },
        ],
    },
    {
        "name": "Krushna Mane",
        "enrolment": "04",
        "role": "Testing, Deployment & Documentation",
        "weeks": [
            {
                "week": 1,
                "rows": [
                    ("Monday",    "08/12/25", "Review of problem statement and project objective", "Clear understanding of MANDIMITRA quality and delivery goals"),
                    ("Tuesday",   "09/12/25", "Study of software testing methodologies for ML-based projects", "Unit, integration, and system testing approaches understood"),
                    ("Wednesday", "10/12/25", "Literature survey on CI/CD and deployment for Python applications", "Knowledge of Docker, Render, and cloud deployment gained"),
                    ("Thursday",  "11/12/25", "Analysis of documentation standards for MSBTE capstone projects", "Report format, chapter structure, and submission norms identified"),
                    ("Friday",    "12/12/25", "Environment setup — pytest, Docker Desktop, Git, documentation tools", "Testing and deployment tools installed and configured"),
                ],
            },
            {
                "week": 2,
                "rows": [
                    ("Monday",    "15/12/25", "Design of test plan for MANDIMITRA — scope, categories, priorities", "Comprehensive test plan document with 50+ test cases drafted"),
                    ("Tuesday",   "16/12/25", "Study of project.yaml and data_sources.yaml for config validation", "Configuration file schemas and valid ranges documented"),
                    ("Wednesday", "17/12/25", "Setup of project Git repository — branching strategy and .gitignore", "Git workflow with feature branches and review process established"),
                    ("Thursday",  "18/12/25", "Initial README.md structure — sections, badges, and quick start guide", "README skeleton with all required sections outlined"),
                    ("Friday",    "19/12/25", "Research on MSBTE capstone report format — 316004 guidelines", "Report structure with front pages, chapters, and appendix planned"),
                ],
            },
            {
                "week": 3,
                "rows": [
                    ("Monday",    "22/12/25", "Unit test development for data pipeline — download, validate, merge", "15 unit tests for pipeline functions written and passing"),
                    ("Tuesday",   "23/12/25", "Unit test development for district normalization and dedup logic", "Edge cases — misspelled districts, duplicates — covered"),
                    ("Wednesday", "24/12/25", "Schema validation tests using Pandera — valid and invalid data", "10 schema tests ensuring Maharashtra-only constraint works"),
                    ("Thursday",  "25/12/25", "API endpoint tests — health check, authentication, error responses", "pytest-based API tests for all FastAPI routes created"),
                    ("Friday",    "26/12/25", "Test automation script — single command to run all test suites", "Makefile/script running all tests with coverage report"),
                ],
            },
            {
                "week": 4,
                "rows": [
                    ("Monday",    "29/12/25", "Integration testing — pipeline → database → API flow", "End-to-end data flow validated from download to API response"),
                    ("Tuesday",   "30/12/25", "ML model inference tests — input validation, output format checks", "All 3 model endpoints returning correct JSON structure"),
                    ("Wednesday", "31/12/25", "Performance benchmarking — API response times under load", "Baseline latency metrics recorded for all endpoints"),
                    ("Thursday",  "01/01/26", "Documentation — Chapter 1 (Introduction) draft for capstone report", "Introduction with problem background and objectives written"),
                    ("Friday",    "02/01/26", "Documentation — Chapter 2 (Literature Survey) draft", "5 research papers summarized with technique adoption mapping"),
                ],
            },
            {
                "week": 5,
                "rows": [
                    ("Monday",    "05/01/26", "Documentation — Chapter 3 (Scope of the Project) draft", "Project scope, limitations, and future expansion documented"),
                    ("Tuesday",   "06/01/26", "Documentation — Chapter 4 (Methodology/Approach) draft", "Pipeline architecture, ML approach, and API design explained"),
                    ("Wednesday", "07/01/26", "Creation of architecture.png and pipeline_flowchart.png diagrams", "System architecture and data flow diagrams designed"),
                    ("Thursday",  "08/01/26", "Creation of DFD Level-0 and module interaction diagrams", "dfd_level0.png and module_interaction.png designed"),
                    ("Friday",    "09/01/26", "Validation script development — verify production standards", "Self-check script testing all quality constraints"),
                ],
            },
            {
                "week": 6,
                "rows": [
                    ("Monday",    "12/01/26", "Docker containerization — writing Dockerfile for backend", "Multi-stage Dockerfile with Python 3.10 slim image"),
                    ("Tuesday",   "13/01/26", "Docker testing — build, run, and verify container locally", "Container running with all API endpoints functional"),
                    ("Wednesday", "14/01/26", "Docker Compose setup for local development (API + DB)", "docker-compose.yml with FastAPI and PostgreSQL services"),
                    ("Thursday",  "15/01/26", "Render deployment configuration using render.yaml", "Build command, start command, and env vars configured"),
                    ("Friday",    "16/01/26", "Initial deployment to Render — troubleshoot and verify", "Backend live on Render with successful health check"),
                ],
            },
            {
                "week": 7,
                "rows": [
                    ("Monday",    "19/01/26", "Documentation — Chapter 5 (Design, Working, Processes) draft", "Detailed system design with all module descriptions"),
                    ("Tuesday",   "20/01/26", "Creation of testing_validation_flow.png and dedup_flowchart.png", "Testing workflow and data deduplication diagrams designed"),
                    ("Wednesday", "21/01/26", "Documentation — Chapter 6 (Results and Applications) draft", "Model results, data coverage, and application screenshots"),
                    ("Thursday",  "22/01/26", "Creation of data_coverage.png and scope_diagram.png", "Data coverage visualization and project scope diagram"),
                    ("Friday",    "23/01/26", "Documentation — Conclusion and References sections", "Conclusion with future work and 15 properly cited references"),
                ],
            },
            {
                "week": 8,
                "rows": [
                    ("Monday",    "26/01/26", "Regression testing after backend optimization changes", "All 50+ test cases passing after code changes"),
                    ("Tuesday",   "27/01/26", "Security testing — SQL injection, XSS, and auth bypass checks", "No security vulnerabilities found in API endpoints"),
                    ("Wednesday", "28/01/26", "Data quality audit — generating completeness and validity reports", "data_completeness reports generated for all datasets"),
                    ("Thursday",  "29/01/26", "Data quality optimization — fixing issues found in audit", "data_quality_optimization_report.md with all fixes documented"),
                    ("Friday",    "30/01/26", "Cross-platform testing — API calls from web, mobile, and curl", "Consistent API behavior across all client types verified"),
                ],
            },
            {
                "week": 9,
                "rows": [
                    ("Monday",    "02/02/26", "Report generation script — generate_report.py development", "Python script using ReportLab to generate formatted PDF"),
                    ("Tuesday",   "03/02/26", "Report script — cover page, certificate, acknowledgement pages", "Front pages matching MSBTE format with styling"),
                    ("Wednesday", "04/02/26", "Report script — embedding all 18 diagrams into chapters", "All diagrams auto-inserted at correct positions"),
                    ("Thursday",  "05/02/26", "Report script — table of contents, list of figures/tables", "Auto-generated TOC, LOF, and LOT pages"),
                    ("Friday",    "06/02/26", "Full PDF report generation and review", "Complete capstone report PDF generated successfully"),
                ],
            },
            {
                "week": 10,
                "rows": [
                    ("Monday",    "09/02/26", "Production deployment verification — all services health checked", "Backend, database, and model serving all operational"),
                    ("Tuesday",   "10/02/26", "BUILD.md documentation — how to regenerate report from source", "Step-by-step build instructions for report PDF"),
                    ("Wednesday", "11/02/26", "CAPSTONE_PPT_GUIDE.md — 13-slide presentation structure", "Complete PPT guide with speaker notes for all slides"),
                    ("Thursday",  "12/02/26", "Final test run — complete regression suite on production", "All tests green on production deployment"),
                    ("Friday",    "13/02/26", "Bug tracking and resolution — filing and fixing remaining issues", "All known bugs resolved, zero open issues"),
                ],
            },
            {
                "week": 11,
                "rows": [
                    ("Monday",    "16/02/26", "Report proofreading — grammar, formatting, and content review", "All chapters reviewed and corrections applied"),
                    ("Tuesday",   "17/02/26", "Diagram polishing — ensuring all 18 diagrams are print-quality", "High-resolution diagrams with consistent styling"),
                    ("Wednesday", "18/02/26", "PPT creation from CAPSTONE_PPT_GUIDE — visual slides with diagrams", "13-slide PowerPoint presentation created"),
                    ("Thursday",  "19/02/26", "Capstone daily diary compilation for all 4 team members", "12-week daily diaries for all students formatted"),
                    ("Friday",    "20/02/26", "Final report PDF generation with all corrections and updates", "Publication-ready capstone report PDF finalized"),
                ],
            },
            {
                "week": 12,
                "rows": [
                    ("Monday",    "23/02/26", "Final system testing — complete end-to-end validation", "All modules integrated and functioning correctly"),
                    ("Tuesday",   "24/02/26", "Submission package preparation — report, PPT, code, APK, diary", "All deliverables organized in submission folder"),
                    ("Wednesday", "25/02/26", "Demo rehearsal — testing live demo on projector setup", "Demo flow verified on presentation hardware"),
                    ("Thursday",  "26/02/26", "Practice presentation — full team dry run with Q&A", "All team members confident with their sections"),
                    ("Friday",    "27/02/26", "Final submission — all documents, code, and artifacts handed over", "Complete capstone project submitted and archived"),
                ],
            },
        ],
    },
]

GUIDE_NAME  = "Mayur Gund"
SEMESTER    = "AN-6-K"


# ── Helpers ───────────────────────────────────────────────────

def set_cell_border(cell, **kwargs):
    """Set cell borders. Usage: set_cell_border(cell, top={"sz":4, "val":"single"}, ...)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            element = OxmlElement(f"w:{edge}")
            for attr in ("sz", "val", "color", "space"):
                if attr in kwargs[edge]:
                    element.set(qn(f"w:{attr}"), str(kwargs[edge][attr]))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """Set borders for every cell in the table."""
    border = {"sz": 4, "val": "single", "color": "000000", "space": "0"}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border)


def add_header_info(doc, student):
    """Add the header section with student info."""
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("CAPSTONE PROJECT DAILY DIARY")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    # Student details
    fields = [
        ("Name of the Student: –", student["name"]),
        ("Name of Guide (Faculty) :–", GUIDE_NAME),
        ("Enrolment Number: –", student["enrolment"]),
        ("Semester:–", SEMESTER),
        ("Role Focus :", student["role"]),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run_label = p.add_run(label + " ")
        run_label.bold = True
        run_label.font.size = Pt(11)
        run_value = p.add_run(value)
        run_value.font.size = Pt(11)


def add_week(doc, week_data, is_first_week=False):
    """Add one week table to the document."""
    week_num = week_data["week"]

    doc.add_paragraph()  # spacer

    # Week heading
    heading = doc.add_paragraph()
    run = heading.add_run(f"WEEK: {week_num}")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # Table: 5 columns
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Column widths
    widths = [Cm(2.5), Cm(2.2), Cm(6.5), Cm(5.5), Cm(2.3)]

    # Header row
    headers = ["Day", "Date", "Activity Carried Out", "Achievement of milestone/step as per plan", "Remark of Faculty"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for day, date, activity, achievement in week_data["rows"]:
        row = table.add_row()
        vals = [day, date, activity, achievement, ""]
        for i, val in enumerate(vals):
            row.cells[i].width = widths[i]
            p = row.cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(9)

    set_table_borders(table)

    # Add spacer paragraphs to push signature toward page bottom
    # First page (with header) has less space, subsequent pages have more
    spacer_count = 4 if is_first_week else 8
    for _ in range(spacer_count):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)

    # Signature line at bottom
    sig = doc.add_paragraph()
    sig.paragraph_format.space_before = Pt(6)
    sig.paragraph_format.space_after = Pt(0)
    run1 = sig.add_run("Dated Signature of Faculty")
    run1.bold = True
    run1.font.size = Pt(11)
    run2 = sig.add_run("                                                    ")
    run2.font.size = Pt(11)
    run3 = sig.add_run("Dated Signature of HOD")
    run3.bold = True
    run3.font.size = Pt(11)


def generate_diary(student):
    """Generate a .docx file for one student."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    add_header_info(doc, student)

    for i, week_data in enumerate(student["weeks"]):
        is_first = (i == 0)
        add_week(doc, week_data, is_first_week=is_first)
        # Page break after every week (except the last one)
        if week_data["week"] < 12:
            doc.add_page_break()

    filename = student["name"].replace(" ", "_") + "_Daily_Diary.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    print(f"  ✅ Generated: {filepath}")
    return filepath


# ── Main ──────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating Daily Diary .docx files...")
    for student in STUDENTS:
        generate_diary(student)
    print("\nAll 4 diary files generated successfully!")
