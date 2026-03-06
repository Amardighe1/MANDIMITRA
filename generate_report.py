"""
Generate MandiMitra Project Report (DOCX + PDF)
Follows the reference report format but with real MandiMitra project content.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from fpdf import FPDF
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ═══════════════════════════════════════════════════════════════════
#  REPORT CONTENT - MandiMitra
# ═══════════════════════════════════════════════════════════════════

TITLE = "MandiMitra - AI-Powered Agricultural Intelligence Platform"

SECTIONS = [
    # ── 1. Introduction ──
    {
        "heading": "1. Introduction",
        "paragraphs": [
            (
                'MandiMitra ("Market Friend") is an AI-powered agricultural intelligence platform '
                "built exclusively for Maharashtra state farmers. The project combines real-time "
                "mandi (market) price data, weather forecasts, ML-based crop risk assessment, "
                "price forecasting, crop disease detection, a buyer marketplace, and veterinary "
                "services into a single mobile and web application."
            ),
            (
                "The platform is designed to empower Maharashtra's farming community by providing "
                "data-driven insights that help farmers make informed decisions about when to sell "
                "their crops, assess weather-related risks, detect crop diseases early, and access "
                "veterinary services in emergencies. MandiMitra covers all 36 districts of "
                "Maharashtra with pre-configured district headquarters coordinates."
            ),
        ],
    },
    # ── 2. Project Objectives ──
    {
        "heading": "2. Project Objectives",
        "bullets": [
            "Provide real-time mandi price intelligence from AGMARKNET for Maharashtra markets.",
            "Build ML models for 15-day crop price forecasting with confidence intervals.",
            "Develop a hybrid rules + ML crop risk advisory system using weather and crop lifecycle data.",
            "Implement on-device crop disease detection using deep learning (17 diseases across 5 crops).",
            "Create a buyer marketplace where buyers post daily prices and farmers compare them.",
            "Offer veterinary services with emergency SOS and auto-assignment to the nearest verified doctor.",
            "Deliver a production-ready mobile app (Android APK) and web dashboard for farmers.",
        ],
    },
    # ── 3. System Architecture ──
    {
        "heading": "3. System Architecture",
        "subsections": [
            {
                "subheading": "3.1 Project Directory Structure",
                "body": "The project is organized as follows:",
                "code": (
                    "mandimitra/\n"
                    "  |-- api/              # FastAPI backend (auth, buyer, vet, crop_disease, weather)\n"
                    "  |-- configs/          # crop_lifecycle.json, project.yaml, locations CSV\n"
                    "  |-- Crop Diseases/    # Training images (17 classes, 5 crops)\n"
                    "  |-- data/             # raw/, processed/, metadata/ (parquet, CSV)\n"
                    "  |-- models/           # LightGBM + Keras trained model artifacts\n"
                    "  |-- scripts/          # 26 scripts: download, train, validate, pipeline\n"
                    "  |-- src/              # Shared utilities, schemas, HTTP client\n"
                    "  |-- web/              # Next.js 14 + Capacitor Android app\n"
                    "  |-- Dockerfile        # Production Docker image\n"
                    "  |-- render.yaml       # Render deployment blueprint\n"
                    "  |-- requirements.txt  # Python dependencies"
                ),
            },
            {
                "subheading": "3.2 Technology Stack",
                "sub_bullets_grouped": [
                    ("Frontend", [
                        "Next.js 14 (React 18) with static export and Tailwind CSS",
                        "Capacitor for Android APK packaging (on-device TensorFlow.js inference)",
                        "Leaflet / React-Leaflet for interactive maps",
                        "Recharts for data visualization and Framer Motion for animations",
                    ]),
                    ("Backend", [
                        "FastAPI (Python 3.10+) with Uvicorn ASGI server",
                        "Motor (async MongoDB driver) for database operations",
                        "Pydantic models for request/response validation",
                        "JWT-based authentication with role-based access (farmer/doctor/admin)",
                    ]),
                    ("Machine Learning", [
                        "LightGBM for crop risk classification and price forecasting",
                        "TensorFlow/Keras (MobileNetV2) for crop disease detection",
                        "Scikit-learn for preprocessing and evaluation",
                        "Google Gemini AI for disease treatment advice generation",
                    ]),
                    ("Database and Deployment", [
                        "MongoDB Atlas (cloud-hosted NoSQL database)",
                        "Docker containerized deployment on Render (free tier)",
                        "Data.gov.in, NASA POWER, Open-Meteo, and Kaggle as data sources",
                    ]),
                ],
            },
        ],
    },
    # ── 4. Methodology ──
    {
        "heading": "4. Methodology",
        "subsections": [
            {
                "subheading": "4.1 Data Collection",
                "paragraphs": [
                    (
                        "MandiMitra uses a robust, resumable data pipeline to collect agricultural data "
                        "from multiple sources. Current mandi prices are fetched daily from the Data.gov.in "
                        "AGMARKNET API. Historical price data (3.5M+ records) is sourced from Kaggle archives. "
                        "Weather data is collected from NASA POWER (10 years of daily precipitation, temperature, "
                        "and humidity for 36 district HQs) and Open-Meteo (16-day forecasts, free, no API key)."
                    ),
                    (
                        "The pipeline supports chunked, by-district downloads with atomic progress tracking, "
                        "adaptive rate limiting via a token-bucket algorithm, and DuckDB-based deduplication "
                        "for memory-efficient processing of 6M+ mandi rows."
                    ),
                ],
            },
            {
                "subheading": "4.2 Model Training",
                "sub_bullets_grouped": [
                    ("Crop Risk Advisor (LightGBM)", [
                        "Trained as a 3-class classifier (Low / Medium / High risk) with class weights 1:10:50.",
                        "Uses weather features + crop lifecycle stage configuration (sowing windows, sensitive periods).",
                        "5-fold TimeSeriesSplit cross-validation: mean accuracy 88.0%, F1-macro 0.62.",
                        "Final accuracy: 87.4%, F1-weighted: 0.886. SHAP feature importance exported.",
                    ]),
                    ("Price Intelligence Engine (LightGBM)", [
                        "Multi-horizon direct forecasting (1, 3, 7, 14, 15 days ahead).",
                        "60 engineered features: price lags, rolling stats, momentum, calendar encoding, weather variables.",
                        "Trained on ~3.3M rows. 1-day MAE: 362, R-squared: 0.933; 7-day R-squared: 0.890.",
                        "Conformal prediction intervals for calibrated uncertainty (q90 = 924 Rs).",
                        "HOLD/SELL recommendations based on forecasted price change vs. 5% threshold.",
                    ]),
                    ("Crop Disease Detector (MobileNetV2)", [
                        "Transfer learning on MobileNetV2 with 2-phase training (freeze then fine-tune).",
                        "17 classes across 5 crops: Corn, Potato, Rice, Sugarcane, Wheat.",
                        "Validation accuracy: 93.5% on 224x224 RGB images.",
                        "Deployed on-device via TensorFlow.js in the Android APK (works offline).",
                    ]),
                ],
            },
            {
                "subheading": "4.3 Implementation",
                "bullets": [
                    "Farmers access the platform via an Android app (Capacitor APK) or a web dashboard.",
                    "The FastAPI backend serves ML predictions, live prices, weather data, and marketplace features.",
                    "Crop disease detection runs on-device using TF.js, saving ~400 MB of server resources.",
                    "Google Gemini AI generates structured treatment advice for detected diseases.",
                    "The buyer marketplace supports Marathi + English crop names for bilingual accessibility.",
                    "Emergency SOS uses Haversine distance to auto-assign the nearest verified veterinary doctor.",
                    "Background asyncio tasks auto-refresh mandi prices from Data.gov.in every 24 hours.",
                ],
            },
        ],
    },
    # ── 5. Key Features ──
    {
        "heading": "5. Key Features",
        "subsections": [
            {
                "subheading": "5.1 Mandi Price Intelligence",
                "body": (
                    "Real-time daily commodity prices from Data.gov.in AGMARKNET, auto-refreshed every "
                    "24 hours with filtering, sorting, and pagination. Supports all major Maharashtra "
                    "agricultural markets and commodities."
                ),
            },
            {
                "subheading": "5.2 Price Forecasting",
                "body": (
                    "15-day price predictions using LightGBM with 5 forecast horizons (1/3/7/14/15 days). "
                    "Includes statistically calibrated conformal prediction intervals and actionable "
                    "HOLD/SELL recommendations based on a 5% price-change threshold."
                ),
            },
            {
                "subheading": "5.3 Crop Risk Advisory",
                "body": (
                    "Hybrid rules + ML system that classifies crop risk as Low, Medium, or High using "
                    "16-day weather forecasts and crop-specific lifecycle stages. Supports multiple "
                    "crops including Soybean, Cotton, Tur, and Urad with Kharif/Rabi classification."
                ),
            },
            {
                "subheading": "5.4 Crop Disease Detection",
                "body": (
                    "CNN-based (MobileNetV2) image classification for 17 diseases across 5 crops "
                    "(Corn, Potato, Rice, Sugarcane, Wheat) with 93.5% validation accuracy. Runs "
                    "entirely on-device via TensorFlow.js for offline capability. Google Gemini AI "
                    "provides treatment advice."
                ),
            },
            {
                "subheading": "5.5 Veterinary Services and Emergency SOS",
                "body": (
                    "Full veterinary booking system with doctor verification workflow, appointment "
                    "scheduling, and emergency SOS. The SOS system automatically assigns the nearest "
                    "verified doctor using Haversine distance calculation with an escalation chain."
                ),
            },
            {
                "subheading": "5.6 Buyer Marketplace",
                "body": (
                    "Buyers post daily crop prices at their mandi; farmers search and compare best "
                    "prices across markets. Supports 20 predefined crops with both Marathi and "
                    "English names for bilingual accessibility."
                ),
            },
        ],
    },
    # ── 6. Challenges and Solutions ──
    {
        "heading": "6. Challenges and Solutions",
        "subsections": [
            {
                "subheading": "6.1 Data Quality and Scale",
                "body": (
                    "Challenge: Processing 3.5M+ historical mandi records with duplicates and missing values.\n"
                    "Solution: DuckDB-based deduplication with priority rules and Pandera schema validation "
                    "with strict Maharashtra-only checks."
                ),
            },
            {
                "subheading": "6.2 Model Performance",
                "body": (
                    "Challenge: Achieving reliable multi-horizon price forecasts and risk classification.\n"
                    "Solution: 60 engineered features, conformal prediction intervals for uncertainty "
                    "quantification, SHAP explainability for the crop risk model, and class weighting "
                    "(1:10:50) to handle imbalanced risk classes."
                ),
            },
            {
                "subheading": "6.3 Deployment Constraints",
                "body": (
                    "Challenge: Deploying ML models within Render free-tier resource limits.\n"
                    "Solution: Shipping only LightGBM models (~24 MB) on the server, running "
                    "TensorFlow.js disease detection on-device in the APK (saving ~400 MB), "
                    "and Docker multi-stage optimization."
                ),
            },
            {
                "subheading": "6.4 API Rate Limiting",
                "body": (
                    "Challenge: Data.gov.in API rate limits and intermittent failures during bulk downloads.\n"
                    "Solution: Token-bucket rate limiting, exponential backoff, resumable chunked "
                    "downloads with atomic progress files, and ThreadPoolExecutor for parallel requests."
                ),
            },
            {
                "subheading": "6.5 Offline Accessibility",
                "body": (
                    "Challenge: Farmers in remote areas with limited internet connectivity.\n"
                    "Solution: On-device TF.js inference for crop disease detection works fully "
                    "offline once the app is installed, with no server dependency."
                ),
            },
        ],
    },
    # ── 7. Future Enhancements ──
    {
        "heading": "7. Future Enhancements",
        "bullets": [
            "Expand to additional states beyond Maharashtra with configurable location support.",
            "Add a conversational AI chatbot for symptom-based crop diagnosis and farming advice.",
            "Integrate soil health data and satellite imagery for precision agriculture insights.",
            "Add multilingual voice-based interface for farmers with limited literacy.",
            "Implement push notifications for price alerts, weather warnings, and SOS updates.",
            "Build a community forum for peer-to-peer knowledge sharing among farmers.",
        ],
    },
    # ── 8. Conclusion ──
    {
        "heading": "8. Conclusion",
        "paragraphs": [
            (
                "MandiMitra aims to bridge the gap between agricultural data and actionable intelligence "
                "for Maharashtra's farming community. By combining real-time mandi prices, ML-powered "
                "price forecasting, crop risk advisory, disease detection, and veterinary services into "
                "a single platform, the system empowers farmers to make data-driven decisions."
            ),
            (
                "With production-grade ML models trained on 3.5M+ real records, on-device deep learning "
                "for offline disease detection, and a bilingual buyer marketplace, MandiMitra demonstrates "
                "how AI and modern software engineering can directly impact rural livelihoods. The platform "
                "is deployed and accessible as both a web application and an Android APK, serving all 36 "
                "districts of Maharashtra."
            ),
        ],
    },
]


# ═══════════════════════════════════════════════════════════════════
#  DOCX Generation
# ═══════════════════════════════════════════════════════════════════

BLUE = RGBColor(0x1F, 0x6F, 0xB5)


def _set_font(run, name="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def generate_docx(filepath: str):
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # ── Title ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    _set_font(run, size=20, bold=True, color=BLUE)
    doc.add_paragraph()

    # ── Helpers ──
    def h1(text):
        h = doc.add_heading(level=1)
        run = h.add_run(text)
        _set_font(run, size=16, bold=True, color=BLUE)

    def h2(text):
        h = doc.add_heading(level=2)
        run = h.add_run(text)
        _set_font(run, size=14, bold=True, color=BLUE)

    def para(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_font(run)
        p.paragraph_format.space_after = Pt(6)

    def bullets(items):
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            if p.runs:
                for r in p.runs:
                    r.text = ""
            run = p.add_run(item)
            _set_font(run)

    def code_block(text):
        for line in text.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(line)
            _set_font(run, name="Courier New", size=10)

    def grouped_bullets(groups):
        for label, items in groups:
            p = doc.add_paragraph()
            run = p.add_run(label + ":")
            _set_font(run, bold=True)
            for item in items:
                bp = doc.add_paragraph(style="List Bullet")
                if bp.runs:
                    for r in bp.runs:
                        r.text = ""
                run = bp.add_run(item)
                _set_font(run)

    # ── Render sections ──
    for sec in SECTIONS:
        h1(sec["heading"])

        for text in sec.get("paragraphs", []):
            para(text)

        if "body" in sec and sec["body"]:
            for line in sec["body"].split("\n"):
                para(line)

        if "bullets" in sec:
            bullets(sec["bullets"])

        for sub in sec.get("subsections", []):
            h2(sub["subheading"])
            for text in sub.get("paragraphs", []):
                para(text)
            if sub.get("body"):
                for line in sub["body"].split("\n"):
                    para(line)
            if "bullets" in sub:
                bullets(sub["bullets"])
            if "sub_bullets_grouped" in sub:
                grouped_bullets(sub["sub_bullets_grouped"])
            if "code" in sub:
                code_block(sub["code"])

    doc.save(filepath)
    print(f"DOCX saved -> {filepath}")


# ═══════════════════════════════════════════════════════════════════
#  PDF Generation
# ═══════════════════════════════════════════════════════════════════

class ReportPDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font("Times", "I", 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 6, "MandiMitra - AI-Powered Agricultural Intelligence Platform", align="C")
            self.ln(8)

    def footer(self):
        self.set_y(-15)
        self.set_font("Times", "I", 8)
        self.set_text_color(120, 120, 120)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def generate_pdf(filepath: str):
    pdf = ReportPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin

    def title(text):
        pdf.set_font("Times", "B", 20)
        pdf.set_text_color(31, 111, 181)
        pdf.multi_cell(W, 12, text, align="C")
        pdf.ln(6)

    def h1(text):
        pdf.ln(4)
        pdf.set_font("Times", "B", 16)
        pdf.set_text_color(31, 111, 181)
        pdf.cell(W, 10, text)
        pdf.ln(11)

    def h2(text):
        pdf.ln(2)
        pdf.set_font("Times", "B", 13)
        pdf.set_text_color(31, 111, 181)
        pdf.cell(W, 8, text)
        pdf.ln(9)

    def body(text):
        pdf.set_font("Times", "", 12)
        pdf.set_text_color(0, 0, 0)
        for line in text.split("\n"):
            pdf.multi_cell(W, 7, line)
            pdf.ln(1)

    def bullet_list(items):
        pdf.set_font("Times", "", 12)
        pdf.set_text_color(0, 0, 0)
        for item in items:
            pdf.cell(8)
            pdf.cell(4, 7, "-")
            pdf.multi_cell(W - 12, 7, " " + item)
            pdf.ln(1)

    def code_block(text):
        pdf.set_font("Courier", "", 9)
        pdf.set_text_color(60, 60, 60)
        for line in text.split("\n"):
            pdf.cell(12)
            pdf.cell(W - 12, 5, line)
            pdf.ln(5)
        pdf.ln(2)

    def grouped_bullets(groups):
        for label, items in groups:
            pdf.set_font("Times", "B", 12)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(W, 7, label + ":")
            pdf.ln(8)
            bullet_list(items)

    # ── Render ──
    title(TITLE)

    for sec in SECTIONS:
        h1(sec["heading"])

        for text in sec.get("paragraphs", []):
            body(text)
        if "body" in sec and sec["body"]:
            body(sec["body"])
        if "bullets" in sec:
            bullet_list(sec["bullets"])

        for sub in sec.get("subsections", []):
            h2(sub["subheading"])
            for text in sub.get("paragraphs", []):
                body(text)
            if sub.get("body"):
                body(sub["body"])
            if "bullets" in sub:
                bullet_list(sub["bullets"])
            if "sub_bullets_grouped" in sub:
                grouped_bullets(sub["sub_bullets_grouped"])
            if "code" in sub:
                code_block(sub["code"])

    pdf.output(filepath)
    print(f"PDF saved  -> {filepath}")


# ═══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    docx_path = os.path.join(OUTPUT_DIR, "MandiMitra_Project_Report.docx")
    pdf_path = os.path.join(OUTPUT_DIR, "MandiMitra_Project_Report.pdf")
    generate_docx(docx_path)
    generate_pdf(pdf_path)
    print("\nDone! Both files generated.")
